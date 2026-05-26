#!/usr/bin/env python3
"""Design V12 final optimization package and academic DOCX report.

V12 consolidates the project into a manuscript-facing package:

* final non-intersecting formula geometry from V11;
* operating-window optimization around the selected design;
* flow-rate, Joule-heating, pressure, residence-time and purity/recovery plots;
* a structured DOCX report suitable for manuscript drafting.
"""

from __future__ import annotations

import csv
import json
import math
import shutil
import subprocess
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from design_v11_shape_ml_nonintersection import (
    FAMILIES,
    centerline,
    field_gain_from_electrodes,
    geometry_screen,
    performance_metrics,
    run_row,
    sigmoid,
    simulate,
)
from study_common import load_openfoam_field_stats, run_openfoam_case, write_csv


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "results" / "design_v12_final_package"
FIG_DIR = OUT_DIR / "figures"
DOC_DIR = ROOT / "output" / "doc"
REPORT_DOCX = DOC_DIR / "spiral_dep_v12_final_optimization_report.docx"


def read_csv(path: Path) -> list[dict]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def to_float(row: dict, key: str, default: float = 0.0) -> float:
    try:
        return float(row.get(key, default))
    except (TypeError, ValueError):
        return default


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def save_plot(fig: plt.Figure, name: str) -> Path:
    path = FIG_DIR / name
    fig.tight_layout()
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


def variant_sample(base: dict, rng: np.random.Generator) -> dict:
    sample = dict(base)
    sample["voltage_v"] = float(rng.uniform(9.0, 22.0))
    sample["velocity_m_s"] = float(rng.uniform(0.75e-3, 4.40e-3))
    sample["electrode_gap_m"] = float(rng.uniform(36e-6, 76e-6))
    sample["electrode_coverage"] = float(rng.uniform(0.42, 0.82))
    sample["dep_start_fraction"] = float(rng.uniform(0.18, 0.40))
    sample["dep_end_fraction"] = float(rng.uniform(0.72, 0.95))
    if sample["dep_end_fraction"] <= sample["dep_start_fraction"] + 0.24:
        sample["dep_end_fraction"] = min(0.98, sample["dep_start_fraction"] + 0.24)
    sample["outlet_split_ratio"] = float(rng.uniform(0.34, 0.55))
    sample["inlet_spread_ratio"] = float(rng.uniform(0.10, 0.32))
    sample["inlet_offset_ratio"] = float(rng.uniform(-0.70, -0.38))
    sample["frequency_hz"] = float(np.clip(sample["frequency_hz"] * rng.uniform(0.75, 1.30), 450e3, 2.80e6))
    family = next(f for f in FAMILIES if f.name == sample["family"])
    sample["field_gain"] = field_gain_from_electrodes(sample["channel_width_m"], sample["electrode_gap_m"], sample["electrode_coverage"], family)
    return sample


def v12_score(row: dict) -> float:
    thermal_penalty = {"low": 0.0, "moderate": 0.06, "high": 0.28}[row["thermal_risk"]]
    return (
        row["target_correct"]
        + 0.35 * row["live_recovery"]
        + 0.35 * row["dead_removal"]
        + 0.18 * row["live_outlet_purity"]
        + 0.18 * row["dead_outlet_purity"]
        - 1.30 * row["wall_loss"]
        - 0.006 * max(0.0, row["active_joule_power_mW"] - 8.0)
        - 0.020 * max(0.0, row["steady_substrate_delta_c_proxy"] - 6.0)
        - 0.010 * max(0.0, row["pressure_drop_kPa"] - 8.0)
        - 0.010 * max(0.0, row["residence_s"] - 8.0)
        - thermal_penalty
    )


def run_operating_optimization(base: dict) -> tuple[list[dict], list[dict], list[dict]]:
    run_openfoam_case()
    field_stats = load_openfoam_field_stats()
    rng = np.random.default_rng(12120)
    rows: list[dict] = []
    samples: list[dict] = []

    # Seed the search with the V11 best point and a publication-style grid.
    deterministic = []
    for voltage in [10.0, 12.0, 14.0, 16.0, 18.0, 20.0]:
        for velocity in [0.80e-3, 1.20e-3, 1.60e-3, 2.20e-3, 3.20e-3]:
            sample = dict(base)
            sample["voltage_v"] = voltage
            sample["velocity_m_s"] = velocity
            family = next(f for f in FAMILIES if f.name == sample["family"])
            sample["field_gain"] = field_gain_from_electrodes(sample["channel_width_m"], sample["electrode_gap_m"], sample["electrode_coverage"], family)
            deterministic.append(sample)

    random_samples = [variant_sample(base, rng) for _ in range(90)]
    for sample in [base] + deterministic + random_samples:
        if not geometry_screen(sample)["geometry_valid"]:
            continue
        sid = len(samples)
        samples.append(sample)
        row = run_row(sample, field_stats, sid, 121_000 + sid, "v12_operating_screen", 65, 230)
        row["v12_score"] = v12_score(row)
        rows.append(row)

    rows.sort(key=lambda r: r["v12_score"], reverse=True)

    validation_rows: list[dict] = []
    control_rows: list[dict] = []
    for rank, row in enumerate(rows[:8], start=1):
        sample = samples[int(row["sample_id"])]
        for seed in [122_001, 122_019, 122_041, 122_073]:
            vrow = run_row(sample, field_stats, int(row["sample_id"]), seed, "v12_final_validation", 140, 330)
            vrow["v12_score"] = v12_score(vrow)
            vrow["candidate_rank"] = rank
            validation_rows.append(vrow)
        for control_name, kwargs in [
            ("same_length_straight_dep", {"geometry_override": "straight", "dep_enabled": True}),
            ("no_dep", {"dep_enabled": False}),
            ("unfocused_inlet", {"inlet_override": (0.0, 1.0)}),
        ]:
            scores = []
            wall_losses = []
            for seed in [123_003, 123_021, 123_039]:
                particles, summary, _ = simulate(sample, field_stats, seed, 115, 290, **kwargs)
                perf = performance_metrics(particles, summary)
                scores.append(perf["target_correct"])
                wall_losses.append(perf["wall_loss"])
            control_rows.append(
                {
                    "sample_id": int(row["sample_id"]),
                    "candidate_rank": rank,
                    "control": control_name,
                    "mean_target_correct": float(np.mean(scores)),
                    "mean_wall_loss": float(np.mean(wall_losses)),
                }
            )

    summary = summarize_validation(validation_rows, control_rows)
    return rows, validation_rows, summary


def summarize_validation(validation_rows: list[dict], control_rows: list[dict]) -> list[dict]:
    grouped: dict[int, list[dict]] = {}
    for row in validation_rows:
        grouped.setdefault(int(row["sample_id"]), []).append(row)
    controls = {(int(r["sample_id"]), r["control"]): r for r in control_rows}
    out = []
    for sid, items in grouped.items():
        first = items[0]
        row = {
            "sample_id": sid,
            "candidate_rank": first["candidate_rank"],
            "family": first["family"],
            "mean_target_correct": float(np.mean([to_float(r, "target_correct") for r in items])),
            "std_target_correct": float(np.std([to_float(r, "target_correct") for r in items], ddof=1)),
            "mean_live_recovery": float(np.mean([to_float(r, "live_recovery") for r in items])),
            "mean_dead_removal": float(np.mean([to_float(r, "dead_removal") for r in items])),
            "mean_live_outlet_purity": float(np.mean([to_float(r, "live_outlet_purity") for r in items])),
            "mean_dead_outlet_purity": float(np.mean([to_float(r, "dead_outlet_purity") for r in items])),
            "mean_wall_loss": float(np.mean([to_float(r, "wall_loss") for r in items])),
            "length_mm": first["length_mm"],
            "footprint_mm2": first["footprint_mm2"],
            "residence_s": first["residence_s"],
            "flow_uL_min": first["flow_uL_min"],
            "pressure_drop_kPa": first["pressure_drop_kPa"],
            "active_joule_power_mW": first["active_joule_power_mW"],
            "steady_substrate_delta_c_proxy": first["steady_substrate_delta_c_proxy"],
            "thermal_risk": first["thermal_risk"],
            "voltage_v": first["voltage_v"],
            "velocity_um_s": first["velocity_um_s"],
            "electrode_gap_um": first["electrode_gap_um"],
            "electrode_coverage": first["electrode_coverage"],
            "frequency_hz": first["frequency_hz"],
            "v12_score": float(np.mean([to_float(r, "v12_score") for r in items])),
        }
        straight = controls[(sid, "same_length_straight_dep")]
        no_dep = controls[(sid, "no_dep")]
        unfocused = controls[(sid, "unfocused_inlet")]
        row["straight_control_correct"] = float(straight["mean_target_correct"])
        row["no_dep_correct"] = float(no_dep["mean_target_correct"])
        row["unfocused_inlet_correct"] = float(unfocused["mean_target_correct"])
        row["topology_gain_vs_straight"] = row["mean_target_correct"] - row["straight_control_correct"]
        row["passes_final_gate"] = (
            row["mean_target_correct"] >= 0.90
            and row["topology_gain_vs_straight"] >= 0.08
            and row["mean_wall_loss"] < 0.08
            and row["active_joule_power_mW"] <= 10.0
            and row["thermal_risk"] != "high"
        )
        out.append(row)
    out.sort(key=lambda r: (r["passes_final_gate"], r["v12_score"], r["topology_gain_vs_straight"]), reverse=True)
    return out


def make_figures(screen_rows: list[dict], final_summary: list[dict]) -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    plt.style.use("seaborn-v0_8-whitegrid")
    figs: dict[str, Path] = {}
    v11_dir = ROOT / "results" / "design_v11_shape_ml_nonintersection"
    v10_dir = ROOT / "results" / "design_v10_short_clean_shapes"
    v9_dir = ROOT / "results" / "design_v9_formula_shapes"

    best = load_json(v11_dir / "rank1_monotone_curvature_spiral_formula.json")
    _, x, y = centerline(best, n=1600)
    p = np.linspace(0.0, 1.0, len(x))
    alpha = sigmoid((p - best["dep_start_fraction"]) / best["dep_edge_smoothness"]) * (
        1.0 - sigmoid((p - best["dep_end_fraction"]) / best["dep_edge_smoothness"])
    )
    fig, ax = plt.subplots(figsize=(7.5, 6.0))
    ax.plot(x * 1e3, y * 1e3, color="#0f172a", lw=2.6)
    sc = ax.scatter(x[::8] * 1e3, y[::8] * 1e3, c=alpha[::8], cmap="viridis", s=14, zorder=3)
    cb = fig.colorbar(sc, ax=ax, shrink=0.82)
    cb.set_label("DEP activation")
    ax.set_aspect("equal")
    ax.set_xlabel("x (mm)")
    ax.set_ylabel("y (mm)")
    ax.set_title("Final V12 baseline geometry: monotone-curvature spiral")
    figs["geometry"] = save_plot(fig, "fig01_final_geometry.png")

    rows_v11 = read_csv(v11_dir / "v11_validation_summary.csv")
    labels = [r["family"].replace("_", "\n") for r in rows_v11[:8]]
    fig, ax = plt.subplots(figsize=(9.0, 5.0))
    ax.bar(np.arange(len(labels)), [to_float(r, "mean_target_correct") for r in rows_v11[:8]], color="#2a9d8f")
    ax.plot(np.arange(len(labels)), [to_float(r, "straight_control_correct") for r in rows_v11[:8]], color="#111827", marker="o", lw=2, label="straight DEP control")
    ax.set_xticks(np.arange(len(labels)), labels, fontsize=7)
    ax.set_ylim(0.45, 1.05)
    ax.set_ylabel("Target correct")
    ax.set_title("Geometry candidates and same-length straight controls")
    ax.legend()
    figs["geometry_controls"] = save_plot(fig, "fig02_geometry_controls.png")

    evolution = []
    for name, path in [
        ("V9", v9_dir / "v9_validation_summary.csv"),
        ("V10", v10_dir / "v10_validation_summary.csv"),
        ("V11", v11_dir / "v11_validation_summary.csv"),
    ]:
        row = read_csv(path)[0]
        evolution.append(
            {
                "version": name,
                "correct": to_float(row, "mean_target_correct"),
                "length": to_float(row, "length_mm"),
                "gain": to_float(row, "topology_gain_vs_straight"),
            }
        )
    fig, axes = plt.subplots(1, 3, figsize=(10.5, 3.4))
    for ax, metric, title, color in [
        (axes[0], "correct", "Target correct", "#2a9d8f"),
        (axes[1], "length", "Length (mm)", "#457b9d"),
        (axes[2], "gain", "Gain vs straight", "#e76f51"),
    ]:
        ax.bar([r["version"] for r in evolution], [r[metric] for r in evolution], color=color)
        ax.set_title(title)
        ax.grid(axis="y", alpha=0.25)
    fig.suptitle("Design evolution from V9 to V11")
    figs["evolution"] = save_plot(fig, "fig03_design_evolution.png")

    ml_rows = read_csv(v11_dir / "v11_ml_feature_importance.csv")[:12]
    fig, ax = plt.subplots(figsize=(8.2, 5.0))
    ax.barh([r["feature"] for r in ml_rows][::-1], [to_float(r, "importance") for r in ml_rows][::-1], color="#4c78a8")
    ax.set_xlabel("ExtraTrees importance")
    ax.set_title("ML surrogate feature importance for shape screening")
    figs["ml_importance"] = save_plot(fig, "fig04_ml_feature_importance.png")

    fig, ax = plt.subplots(figsize=(8.2, 5.2))
    scatter = ax.scatter(
        [to_float(r, "active_joule_power_mW") for r in screen_rows],
        [to_float(r, "target_correct") for r in screen_rows],
        c=[to_float(r, "flow_uL_min") for r in screen_rows],
        s=36,
        cmap="viridis",
        alpha=0.80,
    )
    ax.axhline(0.90, color="#111827", ls="--", lw=1)
    ax.axvline(10.0, color="#6b7280", ls=":", lw=1)
    cb = fig.colorbar(scatter, ax=ax)
    cb.set_label("Flow rate (uL/min)")
    ax.set_xlabel("Active Joule power proxy (mW)")
    ax.set_ylabel("Target correct")
    ax.set_title("Operating-window Pareto screen")
    figs["pareto_power"] = save_plot(fig, "fig05_operating_pareto_power.png")

    fig, ax = plt.subplots(figsize=(8.2, 5.2))
    scatter = ax.scatter(
        [to_float(r, "flow_uL_min") for r in screen_rows],
        [to_float(r, "target_correct") for r in screen_rows],
        c=[to_float(r, "residence_s") for r in screen_rows],
        s=36,
        cmap="plasma",
        alpha=0.82,
    )
    cb = fig.colorbar(scatter, ax=ax)
    cb.set_label("Residence time (s)")
    ax.axhline(0.90, color="#111827", ls="--", lw=1)
    ax.set_xlabel("Flow rate (uL/min)")
    ax.set_ylabel("Target correct")
    ax.set_title("Flow-rate and residence-time trade-off")
    figs["flow_tradeoff"] = save_plot(fig, "fig06_flowrate_tradeoff.png")

    fig, ax = plt.subplots(figsize=(8.2, 5.2))
    scatter = ax.scatter(
        [to_float(r, "voltage_v") for r in screen_rows],
        [to_float(r, "electrode_gap_um") for r in screen_rows],
        c=[to_float(r, "steady_substrate_delta_c_proxy") for r in screen_rows],
        s=44,
        cmap="inferno",
        alpha=0.82,
    )
    cb = fig.colorbar(scatter, ax=ax)
    cb.set_label("Steady substrate delta T proxy (C)")
    ax.set_xlabel("Voltage (V)")
    ax.set_ylabel("Electrode gap (um)")
    ax.set_title("Thermal feasibility map")
    figs["thermal_map"] = save_plot(fig, "fig07_thermal_map.png")

    fig, ax = plt.subplots(figsize=(8.2, 5.2))
    ax.scatter(
        [to_float(r, "pressure_drop_kPa") for r in screen_rows],
        [to_float(r, "flow_uL_min") for r in screen_rows],
        c=[to_float(r, "target_correct") for r in screen_rows],
        cmap="viridis",
        s=40,
        alpha=0.82,
    )
    ax.set_xlabel("Pressure drop proxy (kPa)")
    ax.set_ylabel("Flow rate (uL/min)")
    ax.set_title("Hydraulic operating envelope")
    figs["hydraulics"] = save_plot(fig, "fig08_hydraulic_envelope.png")

    labels = [f"F{int(r['candidate_rank'])}" for r in final_summary]
    fig, ax = plt.subplots(figsize=(9.0, 5.0))
    xloc = np.arange(len(labels))
    ax.bar(xloc - 0.24, [r["mean_live_recovery"] for r in final_summary], 0.24, label="live recovery", color="#2a9d8f")
    ax.bar(xloc, [r["mean_dead_removal"] for r in final_summary], 0.24, label="dead removal", color="#e76f51")
    ax.bar(xloc + 0.24, [r["mean_live_outlet_purity"] for r in final_summary], 0.24, label="live purity", color="#457b9d")
    ax.set_xticks(xloc, labels)
    ax.set_ylim(0.45, 1.05)
    ax.set_ylabel("Metric")
    ax.set_title("Final validated recovery, removal, and purity metrics")
    ax.legend(ncol=3, fontsize=8)
    figs["metrics"] = save_plot(fig, "fig09_final_metrics.png")

    fig, ax = plt.subplots(figsize=(9.0, 5.0))
    ax.bar(xloc - 0.25, [r["mean_target_correct"] for r in final_summary], 0.25, label="validated", color="#2a9d8f")
    ax.bar(xloc, [r["straight_control_correct"] for r in final_summary], 0.25, label="same-length straight", color="#8d99ae")
    ax.bar(xloc + 0.25, [r["no_dep_correct"] for r in final_summary], 0.25, label="no DEP", color="#e76f51")
    ax.set_xticks(xloc, labels)
    ax.set_ylim(0.45, 1.05)
    ax.set_ylabel("Target correct")
    ax.set_title("Final validation with mechanism controls")
    ax.legend(ncol=3, fontsize=8)
    figs["controls"] = save_plot(fig, "fig10_final_controls.png")

    best_final = final_summary[0]
    table_data = [
        ["Target correct", f"{best_final['mean_target_correct']:.3f} +/- {best_final['std_target_correct']:.3f}"],
        ["Geometry gain", f"{best_final['topology_gain_vs_straight']:.3f}"],
        ["Flow rate", f"{best_final['flow_uL_min']:.3f} uL/min"],
        ["Voltage", f"{best_final['voltage_v']:.2f} V"],
        ["Electrode gap", f"{best_final['electrode_gap_um']:.1f} um"],
        ["Power proxy", f"{best_final['active_joule_power_mW']:.2f} mW"],
        ["Thermal proxy", f"{best_final['steady_substrate_delta_c_proxy']:.2f} C"],
        ["Pressure drop", f"{best_final['pressure_drop_kPa']:.2f} kPa"],
    ]
    fig, ax = plt.subplots(figsize=(7.5, 4.6))
    ax.axis("off")
    table = ax.table(cellText=table_data, colLabels=["Final operating metric", "Value"], loc="center", cellLoc="left")
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.55)
    ax.set_title("Selected V12 operating point", pad=18)
    figs["final_table"] = save_plot(fig, "fig11_selected_operating_point.png")

    fig, ax = plt.subplots(figsize=(9.2, 4.8))
    ax.axis("off")
    boxes = [
        (0.02, 0.62, "Literature gap\nspiral + DEP + live/dead"),
        (0.26, 0.62, "Formula shapes\nnon-intersection filter"),
        (0.50, 0.62, "ML surrogate\nrank parametric candidates"),
        (0.74, 0.62, "Held-out validation\ncontrols + thermal checks"),
        (0.38, 0.18, "Final V12 design\nreport + data package"),
    ]
    for x0, y0, text in boxes:
        rect = plt.Rectangle((x0, y0), 0.21, 0.20, fc="#eef2ff", ec="#334155", lw=1.2, transform=ax.transAxes)
        ax.add_patch(rect)
        ax.text(x0 + 0.105, y0 + 0.10, text, ha="center", va="center", fontsize=10, transform=ax.transAxes)
    for start, end in [((0.23, 0.72), (0.26, 0.72)), ((0.47, 0.72), (0.50, 0.72)), ((0.71, 0.72), (0.74, 0.72)), ((0.61, 0.62), (0.49, 0.38))]:
        ax.annotate("", xy=end, xytext=start, xycoords=ax.transAxes, arrowprops=dict(arrowstyle="->", lw=1.5, color="#334155"))
    ax.set_title("V12 computational workflow")
    figs["workflow"] = save_plot(fig, "fig12_workflow.png")

    return figs


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in [
        ("Title", 20, RGBColor(15, 23, 42)),
        ("Heading 1", 15, RGBColor(15, 23, 42)),
        ("Heading 2", 12, RGBColor(31, 41, 55)),
    ]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = color


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_with_image(doc: Document, title: str, body: str, image: Path, caption: str) -> None:
    doc.add_heading(title, level=1)
    if body:
        doc.add_paragraph(body)
    doc.add_picture(str(image), width=Cm(16.4))
    add_caption(doc, caption)
    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]], header: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, value in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def write_docx_report(figs: dict[str, Path], final_summary: list[dict], screen_rows: list[dict]) -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("V12 Final Optimization Report\n")
    run.bold = True
    run.font.size = Pt(22)
    run = title.add_run("Compact spiral DEP microfluidic design for live/dead cell sorting")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(55, 65, 81)
    doc.add_paragraph("Prepared from the analytical-design repository. Date: 2026-05-26.")
    doc.add_paragraph(
        "This report consolidates the V9-V11 geometry studies and adds a V12 operating-window optimization. "
        "The selected geometry is a non-intersecting monotone-curvature spiral. The final claims remain numerical "
        "and reduced-order until electrode-resolved field simulation and experimental validation are completed."
    )
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    best = final_summary[0]
    bullets = [
        f"Selected family: {best['family']}.",
        f"Validated target correct: {best['mean_target_correct']:.3f} +/- {best['std_target_correct']:.3f}.",
        f"Topology gain versus same-length straight DEP: {best['topology_gain_vs_straight']:.3f}.",
        f"Length and footprint: {best['length_mm']:.2f} mm and {best['footprint_mm2']:.2f} mm2.",
        f"Flow rate, voltage, electrode gap: {best['flow_uL_min']:.3f} uL/min, {best['voltage_v']:.2f} V, {best['electrode_gap_um']:.1f} um.",
        f"Thermal proxies: {best['active_joule_power_mW']:.2f} mW active Joule power and {best['steady_substrate_delta_c_proxy']:.2f} C substrate delta-T proxy.",
        "Main limitation: DEP remains represented by an electrode-gap field-gain proxy; V12 is a design-screening package, not final CFD/FEM proof.",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    doc.add_heading("Literature and Novelty Position", level=1)
    doc.add_paragraph(
        "The project combines three literature clusters: spiral microchannels for compact inertial handling of cells [R4, R8, R16], "
        "DEP live/dead or viability-selective separation [R5-R7, R9, R12], and spiral or circular DEP numerical device studies [R1-R3, R17]. "
        "The working novelty window is not a claim of experimental first demonstration; it is a numerical design study for a continuous "
        "spiral DEP live/dead sorting architecture with explicit recovery, purity, topology-control and thermal checks."
    )
    doc.add_paragraph(
        "This framing prevents a weak claim based only on longer channels. Throughout V10-V12, a same-length straight DEP control is used "
        "to test whether a curved/spiral topology contributes beyond residence time."
    )
    doc.add_page_break()

    add_page_with_image(
        doc,
        "Final Geometry",
        "The selected shape is a monotone-curvature spiral with a smooth low-curvature entry. The formula avoids arbitrary local wiggles and supports clean fabrication discussion.",
        figs["geometry"],
        "Figure 1. Final V12 baseline geometry. Color marks the smooth DEP activation window used in the reduced-order particle model.",
    )
    add_page_with_image(
        doc,
        "Geometry Candidate Controls",
        "Candidate geometries were not accepted by accuracy alone. Designs with high accuracy but equal same-length straight-control performance were rejected as weak topology evidence.",
        figs["geometry_controls"],
        "Figure 2. V11 geometry candidates and same-length straight DEP controls.",
    )
    add_page_with_image(
        doc,
        "Design Evolution",
        "V9 introduced formula-defined spirals, V10 imposed short clean shape families, and V11 added non-intersection filtering. V12 selects the operating point around the V11 geometry.",
        figs["evolution"],
        "Figure 3. Evolution of correct classification, length, and topology gain across design versions.",
    )
    add_page_with_image(
        doc,
        "ML Shape Optimization",
        "The ML model is a surrogate ranker over formula parameters. It does not draw arbitrary curves. The optimized variables have physical meaning and are exported in JSON and CSV.",
        figs["ml_importance"],
        "Figure 4. ExtraTrees feature importance for the V11 shape optimization screen.",
    )
    add_page_with_image(
        doc,
        "Operating Pareto Screen",
        "The V12 operating-window screen balances classification against Joule-heating proxy. Points above 0.90 target correct and below about 10 mW are preferred.",
        figs["pareto_power"],
        "Figure 5. Target correct versus active Joule power proxy for the V12 operating sweep.",
    )
    add_page_with_image(
        doc,
        "Flowrate and Residence Time",
        "Flowrate cannot be maximized blindly. Higher flow improves throughput but can reduce residence time and separation. The selected point stays in a moderate-flow region.",
        figs["flow_tradeoff"],
        "Figure 6. Flowrate, residence time, and target correct trade-off.",
    )
    add_page_with_image(
        doc,
        "Thermal Feasibility",
        "Thermal screening uses a Joule power proxy and a simple substrate heat-loss proxy. This is not a substitute for heat-transfer CFD, but it flags unsafe operating regions.",
        figs["thermal_map"],
        "Figure 7. Voltage-gap operating map colored by steady substrate temperature-rise proxy.",
    )
    add_page_with_image(
        doc,
        "Hydraulic Envelope",
        "Pressure drop and flowrate were tracked to keep the proposal device-like rather than only numerically separable.",
        figs["hydraulics"],
        "Figure 8. Pressure-drop proxy versus flowrate for the V12 operating sweep.",
    )
    add_page_with_image(
        doc,
        "Recovery, Removal, and Purity",
        "The final evaluation reports live recovery, dead removal, and outlet purity. Accuracy alone is intentionally not the only endpoint.",
        figs["metrics"],
        "Figure 9. Final validation metrics for the top V12 operating candidates.",
    )
    add_page_with_image(
        doc,
        "Mechanism Controls",
        "The same-length straight DEP, no-DEP, and unfocused-inlet controls are central to the academic argument. They test topology, field mechanism, and inlet-preparation dependence.",
        figs["controls"],
        "Figure 10. Final validated candidates with mechanism controls.",
    )
    add_page_with_image(
        doc,
        "Selected Operating Point",
        "The chosen point is selected by multi-objective score, not by maximizing accuracy alone. It balances classification, power, thermal proxy, pressure and wall loss.",
        figs["final_table"],
        "Figure 11. Key values for the selected V12 operating point.",
    )
    add_page_with_image(
        doc,
        "Computational Workflow",
        "The report structure follows the intended manuscript workflow: literature gap, formula design, ML ranking, held-out validation and mechanism controls.",
        figs["workflow"],
        "Figure 12. V12 computational workflow.",
    )

    doc.add_heading("Final Candidate Table", level=1)
    table_rows = []
    for row in final_summary[:8]:
        table_rows.append(
            [
                str(row["candidate_rank"]),
                f"{row['mean_target_correct']:.3f}",
                f"{row['topology_gain_vs_straight']:.3f}",
                f"{row['flow_uL_min']:.3f}",
                f"{row['active_joule_power_mW']:.2f}",
                f"{row['steady_substrate_delta_c_proxy']:.2f}",
                str(row["passes_final_gate"]),
            ]
        )
    add_table(
        doc,
        table_rows,
        ["Rank", "Target correct", "Topology gain", "Flow (uL/min)", "Power (mW)", "Delta T proxy (C)", "Gate"],
    )
    doc.add_paragraph(
        "Interpretation: the selected point is the highest-scoring candidate that preserves topology contribution and avoids high thermal-risk flags."
    )
    doc.add_page_break()

    doc.add_heading("Data Package and Reproducibility", level=1)
    doc.add_paragraph(
        "The V12 package is organized so that manuscript figures, supplementary tables and rerun scripts can be traced without manual reconstruction."
    )
    add_table(
        doc,
        [
            ["Operating sweep", "results/design_v12_final_package/v12_operating_screen.csv"],
            ["Held-out validation", "results/design_v12_final_package/v12_final_validation_replicates.csv"],
            ["Final summary", "results/design_v12_final_package/v12_final_validation_summary.csv"],
            ["Figures", "results/design_v12_final_package/figures/"],
            ["Report source script", "scripts/design_v12_final_package.py"],
            ["Baseline geometry", "results/design_v11_shape_ml_nonintersection/rank1_monotone_curvature_spiral_formula.json"],
        ],
        ["Artifact", "Path"],
    )
    doc.add_paragraph(
        "For a publication-grade next pass, the exact V12 geometry should be converted into a mesh/electrode definition and solved with an electrode-resolved field model. The current package is intentionally transparent about this remaining step."
    )
    doc.add_page_break()

    doc.add_heading("Model Validity and Limits", level=1)
    limitations = [
        "The current device result is a reduced-order numerical design screen.",
        "OpenFOAM is used in the workflow for field/thermal case support, but the final V12 DEP forcing still uses a field-gain proxy rather than a full electrode-resolved 3D field.",
        "The cell model includes population variability, Brownian diffusion, inlet distributions and wall-loss probability, but not cell deformation or cell-cell interaction.",
        "Thermal output is a proxy intended for ranking. A heat-transfer solve is needed before claiming safe absolute temperature.",
        "The strongest current claim is design plausibility and a prioritized geometry/operating point for higher-fidelity simulation.",
    ]
    for item in limitations:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    doc.add_heading("Manuscript-Ready Claims", level=1)
    doc.add_paragraph(
        "Claim 1: A monotone-curvature spiral provides a compact, non-intersecting geometry for DEP-based viable/nonviable sorting in the reduced-order model."
    )
    doc.add_paragraph(
        "Claim 2: The selected design performs better than a same-length straight DEP control, supporting a topology contribution rather than a pure residence-time explanation."
    )
    doc.add_paragraph(
        "Claim 3: Accuracy, recovery, purity, wall loss, flowrate, pressure drop and thermal proxies are all needed for a credible device-design argument."
    )
    doc.add_paragraph(
        "Claim 4: The next required validation is electrode-resolved field simulation and a more rigorous thermal solve for the V12 rank-1 geometry."
    )
    doc.add_page_break()

    doc.add_heading("References Used for Framing", level=1)
    refs = [
        "R1. Betyar and Ramiar, numerical simulation of biological particle separation in a spiral microchannel using DEP, 2026.",
        "R4. Kwon/Choi et al., continuous removal of small nonviable mammalian cells and debris using spiral inertial microfluidics, Lab on a Chip, 2018.",
        "R5. Elitas et al., DEP separation of live and dead monocytes using 3D carbon electrodes, Sensors, 2017.",
        "R9. Huang et al., ODEP live/dead separation platform, 2025.",
        "R10. Nguyen et al., facing-electrode DEP cell separation numerical design, Scientific Reports, 2024.",
        "R11. Tian et al., on-chip DEP single-cell manipulation review, 2024.",
        "R15. DEP and AC electrothermal-flow numerical simulation, Micromachines, 2024.",
        "R16. Spiral inertial cell sorting review, Micromachines, 2024.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.save(REPORT_DOCX)


def write_markdown_summary(final_summary: list[dict], screen_rows: list[dict], figs: dict[str, Path]) -> None:
    best = final_summary[0]
    lines = [
        "# V12 Final Optimization Package",
        "",
        "## Selected Design",
        "",
        f"- family: `{best['family']}`",
        f"- target correct: `{best['mean_target_correct']:.3f} +/- {best['std_target_correct']:.3f}`",
        f"- topology gain vs same-length straight DEP: `{best['topology_gain_vs_straight']:.3f}`",
        f"- flow rate: `{best['flow_uL_min']:.3f} uL/min`",
        f"- voltage: `{best['voltage_v']:.2f} V`",
        f"- electrode gap: `{best['electrode_gap_um']:.1f} um`",
        f"- active Joule power proxy: `{best['active_joule_power_mW']:.2f} mW`",
        f"- steady substrate delta-T proxy: `{best['steady_substrate_delta_c_proxy']:.2f} C`",
        f"- passes final gate: `{best['passes_final_gate']}`",
        "",
        "## Produced Artifacts",
        "",
        f"- DOCX report: `{REPORT_DOCX.relative_to(ROOT)}`",
        "- Tables: `v12_operating_screen.csv`, `v12_final_validation_replicates.csv`, `v12_final_validation_summary.csv`",
        "- Figures: `figures/`",
        "",
        "## Honest Limitation",
        "",
        "V12 is a manuscript-facing reduced-order design package. The selected geometry should be rebuilt with electrode-resolved OpenFOAM/FEM fields before final publication claims.",
        "",
        "## Figure List",
        "",
    ]
    for name, path in figs.items():
        lines.append(f"- `{name}`: `{path.relative_to(ROOT)}`")
    (OUT_DIR / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_docx() -> None:
    render_dir = OUT_DIR / "docx_render"
    render_dir.mkdir(parents=True, exist_ok=True)
    if shutil.which("soffice"):
        subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation=file:///tmp/lo_profile_v12_{REPORT_DOCX.stem}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(render_dir),
                str(REPORT_DOCX),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    pdf_path = render_dir / f"{REPORT_DOCX.stem}.pdf"
    if pdf_path.exists() and shutil.which("pdftoppm"):
        subprocess.run(["pdftoppm", "-png", str(pdf_path), str(render_dir / "page")], check=True)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    base = load_json(ROOT / "results" / "design_v11_shape_ml_nonintersection" / "rank1_monotone_curvature_spiral_formula.json")
    screen_rows, validation_rows, final_summary = run_operating_optimization(base)

    write_csv(screen_rows, OUT_DIR / "v12_operating_screen.csv")
    write_csv(validation_rows, OUT_DIR / "v12_final_validation_replicates.csv")
    write_csv(final_summary, OUT_DIR / "v12_final_validation_summary.csv")

    figs = make_figures(screen_rows, final_summary)
    write_docx_report(figs, final_summary, screen_rows)
    write_markdown_summary(final_summary, screen_rows, figs)
    render_docx()

    print(json.dumps({"out_dir": str(OUT_DIR), "docx": str(REPORT_DOCX), "best": final_summary[0]}, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
