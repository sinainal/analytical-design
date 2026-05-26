#!/usr/bin/env python3
"""Literature-bounded feasibility envelope and feature-importance analysis.

This script prepares a publication-style computational package:
1. a literature-derived feasible design envelope,
2. a complete feature catalogue for the current reduced-order model,
3. decision-tree feature importance from existing screening simulations,
4. polished figures and a manuscript-style report.

The analysis is deliberately framed as a computational design/reality-check
stage, not as experimental validation.
"""

from __future__ import annotations

import csv
import json
import math
import shutil
import subprocess
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from sklearn.metrics import mean_absolute_error, r2_score
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeRegressor, export_text, plot_tree


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "results" / "final_v1"
FIG_DIR = OUT_DIR / "figures"
DOC_DIR = ROOT / "output" / "doc"
REPORT_DOCX = DOC_DIR / "feasibility_feature_analysis_report.docx"


INK = "#111827"
MUTED = "#64748b"
BLUE = "#2563eb"
GREEN = "#16a34a"
RED = "#dc2626"
GOLD = "#f59e0b"


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DOC_DIR.mkdir(parents=True, exist_ok=True)


def set_style() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 9,
            "axes.titlesize": 10.5,
            "axes.labelsize": 9,
            "xtick.labelsize": 8,
            "ytick.labelsize": 8,
            "legend.fontsize": 8,
            "axes.linewidth": 0.8,
            "savefig.dpi": 420,
        }
    )


def write_rows(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()), lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


def savefig(fig: plt.Figure, name: str) -> Path:
    path = FIG_DIR / name
    fig.tight_layout()
    fig.savefig(path, dpi=420, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def feasible_ranges() -> list[dict]:
    """Literature-derived feasible envelope for the next design stage."""
    return [
        {
            "parameter": "spiral_turns",
            "proposed_screen_range": "0.75-5.00",
            "preferred_initial_range": "1.00-3.00",
            "units": "turns",
            "literature_anchor": "Betyar and Ramiar varied 1-5 turns; Patel/Xuan used two three-loop spirals.",
            "primary_refs": "R1, R2",
            "design_reason": "Turns must be swept directly because prior spiral DEP work shows voltage decreases with turn count, but extra turns may only add residence time.",
        },
        {
            "parameter": "eccentricity_e",
            "proposed_screen_range": "0.50-1.00",
            "preferred_initial_range": "0.65-1.00",
            "units": "dimensionless",
            "literature_anchor": "Spiral reviews emphasize geometry/cross-section sensitivity; no exact live/dead spiral DEP eccentricity optimum is known.",
            "primary_refs": "R16",
            "design_reason": "Ellipticity is a controlled, formula-defined way to test compactness and curvature distribution without arbitrary curves.",
        },
        {
            "parameter": "channel_width",
            "proposed_screen_range": "50-150",
            "preferred_initial_range": "80-120",
            "units": "um",
            "literature_anchor": "Betyar and Ramiar used 100 um channel width; Patel/Xuan used a 50 um first spiral.",
            "primary_refs": "R1, R2",
            "design_reason": "Width near 100 um is compatible with the spiral DEP precedent and with 20-25 um cells without excessive wall loss.",
        },
        {
            "parameter": "channel_height",
            "proposed_screen_range": "25-100",
            "preferred_initial_range": "40-75",
            "units": "um",
            "literature_anchor": "Nguyen et al. swept 25, 50, 75, and 100 um and reported 50 um as feasible in their DEP device.",
            "primary_refs": "R10",
            "design_reason": "Height changes DEP field exposure and flow resistance; 40-75 um is a conservative first range.",
        },
        {
            "parameter": "inlet_velocity",
            "proposed_screen_range": "800-5000",
            "preferred_initial_range": "1000-3000",
            "units": "um/s",
            "literature_anchor": "Betyar and Ramiar swept 1000, 3000, and 5000 um/s.",
            "primary_refs": "R1",
            "design_reason": "The model must include velocity because higher inlet velocity increases required voltage and reduces residence time.",
        },
        {
            "parameter": "flow_rate",
            "proposed_screen_range": "0.1-10",
            "preferred_initial_range": "0.2-3.0",
            "units": "uL/min",
            "literature_anchor": "Elitas et al. tested 1-10 uL/min and selected 1 uL/min for dead-cell removal.",
            "primary_refs": "R5",
            "design_reason": "Flow-rate targets should be feasible but not selected so low that the model becomes trivially deterministic.",
        },
        {
            "parameter": "voltage",
            "proposed_screen_range": "4-25",
            "preferred_initial_range": "6-22",
            "units": "V or Vpp depending on electrode convention",
            "literature_anchor": "Elitas et al. used 20 Vpp; Huang ODEP swept 2-8 V; Nguyen swept 5-7 V; Betyar and Ramiar used threshold steps.",
            "primary_refs": "R1, R5, R9, R10",
            "design_reason": "Voltage must be optimized with heating and over-deflection checks rather than maximized.",
        },
        {
            "parameter": "frequency",
            "proposed_screen_range": "50 kHz-3 MHz",
            "preferred_initial_range": "100 kHz-1.5 MHz",
            "units": "Hz",
            "literature_anchor": "Elitas et al. characterized 50 kHz-1 MHz and selected 300 kHz; Nguyen used 1 kHz for nDEP cell sorting.",
            "primary_refs": "R5, R10",
            "design_reason": "Frequency is required because live/dead DEP contrast is viability-dependent and can reverse or collapse outside the useful CM window.",
        },
        {
            "parameter": "electrode_gap",
            "proposed_screen_range": "25-100",
            "preferred_initial_range": "40-80",
            "units": "um",
            "literature_anchor": "Betyar and Ramiar used 100 um channel width with side-wall electrodes; Elitas used 50 um diameter, 100 um high 3D electrodes.",
            "primary_refs": "R1, R5",
            "design_reason": "Gap is a field-strength control and a fabrication constraint; too small increases field and heating risks.",
        },
        {
            "parameter": "medium_conductivity",
            "proposed_screen_range": "0.002-0.055",
            "preferred_initial_range": "0.002-0.020",
            "units": "S/m",
            "literature_anchor": "Elitas used 0.002 S/m medium; Nguyen used 55 mS/m PBS; Jiang et al. warn that high conductivity increases Joule heating/ACET.",
            "primary_refs": "R5, R10, R15",
            "design_reason": "Conductivity should be low in the first DEP model unless thermal/ACET coupling is explicitly included.",
        },
        {
            "parameter": "thermal_gate",
            "proposed_screen_range": "DeltaT proxy <5; power proxy <10",
            "preferred_initial_range": "DeltaT proxy <3; power proxy <5",
            "units": "C, mW",
            "literature_anchor": "Jiang et al. show Joule heating creates temperature gradients and ACET disturbance in DEP separation.",
            "primary_refs": "R15",
            "design_reason": "Thermal constraints are used as gates so the optimizer cannot choose unrealistic voltage/conductivity combinations.",
        },
    ]


def latex_equations() -> str:
    return r"""\section*{Mathematical model used in the feasibility analysis}

\subsection*{Spiral geometry}
\[
r(\theta)=r_0+p\theta,\qquad 0\leq \theta \leq 2\pi N
\]

\[
x(\theta)=r(\theta)\cos\theta,\qquad
y(\theta)=e\,r(\theta)\sin\theta
\]

Here \(N\) is the number of turns and \(e\) is the ellipticity ratio. The circular spiral is recovered when \(e=1\).

\subsection*{DEP force scaling}
\[
\mathbf{F}_{DEP}=2\pi \epsilon_m r_p^3 \operatorname{Re}\{K(\omega)\}\nabla |\mathbf{E}|^2
\]

\[
K(\omega)=\frac{\epsilon_p^\ast-\epsilon_m^\ast}{\epsilon_p^\ast+2\epsilon_m^\ast},
\qquad
\epsilon^\ast=\epsilon-\frac{j\sigma}{\omega}
\]

\subsection*{Low-Re particle balance}
\[
6\pi\eta r_p(\mathbf{u}_p-\mathbf{u}_f)=\mathbf{F}_{DEP}+\mathbf{F}_{Dean}+\mathbf{F}_{noise}
\]

\subsection*{Performance metrics}
\[
\mathrm{SE}_c=\frac{N_{c,\mathrm{target\ outlet}}}{N_{c,\mathrm{injected}}},
\qquad
\mathrm{SP}_c=\frac{N_{c,\mathrm{target\ outlet}}}{N_{\mathrm{all,target\ outlet}}}
\]

\[
\mathrm{TopologyGain}=C_{\mathrm{spiral}}-C_{\mathrm{straight,\ same\ length}}
\]
"""


def feature_catalog() -> list[dict]:
    groups = {
        "geometry": [
            ("turns", "Number of spiral revolutions", "primary design variable", "high"),
            ("eccentricity", "Elliptic scaling y=e r sin(theta)", "primary design variable", "high"),
            ("radius_um", "Initial/reference spiral radius", "geometry scale", "high"),
            ("pitch_um_per_turn", "Radial spacing added per turn", "sets compactness and clearance", "high"),
            ("width_um", "Channel width", "cell-wall clearance and field scale", "high"),
            ("length_mm", "Total channel length", "residence time and overfitting risk", "high"),
            ("footprint_mm2", "Planar device area", "fabrication/compactness", "medium"),
            ("mean_curvature_1_m", "Average centerline curvature", "curved-flow/Dean proxy", "medium"),
            ("p95_curvature_1_m", "High-percentile curvature", "bend manufacturability/risk", "medium"),
            ("min_nonlocal_clearance_um", "Minimum spacing between non-neighbor channel segments", "no-overlap gate", "high"),
        ],
        "electrical": [
            ("voltage_v", "Applied electrode voltage", "DEP strength and thermal risk", "high"),
            ("frequency_hz", "AC frequency", "CM-factor/live-dead contrast", "high"),
            ("electrode_gap_um", "Effective electrode gap", "field strength and fabrication", "high"),
            ("electrode_coverage", "Fraction of path with active DEP", "interaction length", "high"),
            ("field_gain", "Reduced-order correction for electrode layout", "proxy for field concentration", "medium"),
            ("effective_field_kV_m", "Effective electric field", "DEP and Joule proxy", "high"),
            ("active_joule_power_mW", "Joule-power proxy", "thermal feasibility", "high"),
            ("steady_substrate_delta_c_proxy", "Temperature-rise proxy", "thermal gate", "high"),
        ],
        "flow": [
            ("velocity_um_s", "Mean inlet velocity", "residence time and hydrodynamic drag", "high"),
            ("flow_uL_min", "Volumetric flow rate", "throughput", "high"),
            ("residence_s", "Length divided by velocity", "interaction time", "high"),
            ("pressure_drop_kPa", "Pressure-drop proxy", "device operation", "medium"),
            ("reynolds", "Reynolds number", "laminar/inertial regime", "medium"),
            ("dean_scale", "Reduced-order curved-flow drift factor", "spiral contribution", "high"),
        ],
        "inlet_outlet": [
            ("inlet_offset_ratio", "Mean inlet lateral offset", "prefocusing sensitivity", "high"),
            ("inlet_spread_ratio", "Distribution width at inlet", "robustness to sample focusing", "high"),
            ("outlet_split_ratio", "Outlet decision boundary", "purity/recovery trade-off", "high"),
            ("dep_start_fraction", "DEP activation start along path", "staging", "medium"),
            ("dep_end_fraction", "DEP activation end along path", "staging", "medium"),
        ],
        "outputs": [
            ("target_correct", "Correct outlet fraction", "primary but insufficient alone", "high"),
            ("live_recovery", "Live cells collected in live outlet", "biological usefulness", "high"),
            ("dead_removal", "Dead cells rejected/removed", "application target", "high"),
            ("live_outlet_purity", "Live fraction in live outlet", "downstream purity", "high"),
            ("dead_outlet_purity", "Dead fraction in dead outlet", "waste stream purity", "medium"),
            ("wall_loss", "Fraction of particles contacting walls/lost", "device failure mode", "high"),
            ("topology_gain_vs_straight", "Spiral improvement over same-length straight control", "spiral-specific evidence", "high"),
        ],
    }
    rows: list[dict] = []
    for group, items in groups.items():
        for name, definition, role, priority in items:
            rows.append({"group": group, "feature": name, "definition": definition, "role": role, "priority": priority})
    return rows


def load_ml_dataset() -> pd.DataFrame:
    paths = [
        ROOT / "results" / "design_v10_short_clean_shapes" / "v10_screening_rows.csv",
        ROOT / "results" / "design_v11_shape_ml_nonintersection" / "v11_screening_rows.csv",
        ROOT / "results" / "design_v12_final_package" / "v12_operating_screen.csv",
    ]
    frames = []
    for path in paths:
        df = pd.read_csv(path)
        df["source_file"] = path.name
        frames.append(df)
    df = pd.concat(frames, ignore_index=True, sort=False)
    df["source_id"] = pd.factorize(df["source_file"])[0]
    df["family_id"] = pd.factorize(df.get("family", "unknown").astype(str))[0]
    df["shape_class_id"] = pd.factorize(df.get("shape_class", "unknown").astype(str))[0]
    df["inlet_kind_id"] = pd.factorize(df.get("inlet_kind", "unknown").astype(str))[0]
    if "manufacturable_no_intersection_ok" in df:
        df["manufacturable_no_intersection_ok_num"] = df["manufacturable_no_intersection_ok"].astype(str).str.lower().map({"true": 1.0, "false": 0.0}).fillna(1.0)
    else:
        df["manufacturable_no_intersection_ok_num"] = 1.0
    if "thermal_risk" in df:
        risk = {"low": 0.0, "moderate": 1.0, "high": 2.0}
        df["thermal_risk_num"] = df["thermal_risk"].astype(str).str.lower().map(risk).fillna(0.0)
    else:
        df["thermal_risk_num"] = 0.0
    df["design_score"] = (
        df["target_correct"].astype(float)
        - 0.55 * df.get("wall_loss", 0).astype(float)
        - 0.015 * df.get("active_joule_power_mW", 0).astype(float).clip(lower=0)
        - 0.010 * df.get("length_mm", 0).astype(float).clip(lower=0)
        - 0.050 * df["thermal_risk_num"]
    )
    return df


def numeric_features(df: pd.DataFrame) -> list[str]:
    candidates = [
        "source_id",
        "family_id",
        "shape_class_id",
        "inlet_kind_id",
        "frequency_hz",
        "voltage_v",
        "effective_voltage_v",
        "velocity_um_s",
        "turns",
        "radius_um",
        "pitch_um_per_turn",
        "width_um",
        "eccentricity",
        "curvature_power",
        "inlet_length_um",
        "inlet_offset_ratio",
        "inlet_spread_ratio",
        "dep_start_fraction",
        "dep_end_fraction",
        "electrode_gap_um",
        "electrode_coverage",
        "field_gain",
        "dean_scale",
        "length_mm",
        "footprint_mm2",
        "residence_s",
        "flow_uL_min",
        "pressure_drop_kPa",
        "reynolds",
        "effective_field_kV_m",
        "active_joule_power_mW",
        "steady_substrate_delta_c_proxy",
        "mean_curvature_1_m",
        "p95_curvature_1_m",
        "min_curvature_radius_um",
        "smoothness_index",
        "min_nonlocal_clearance_um",
        "manufacturable_no_intersection_ok_num",
        "thermal_risk_num",
    ]
    return [c for c in candidates if c in df.columns]


def train_decision_tree(df: pd.DataFrame, target: str) -> tuple[DecisionTreeRegressor, pd.DataFrame, dict, str]:
    features = numeric_features(df)
    data = df[features + [target]].copy()
    for col in features + [target]:
        data[col] = pd.to_numeric(data[col], errors="coerce")
    data = data.dropna(subset=[target])
    for col in features:
        data[col] = data[col].fillna(data[col].median())
    x = data[features]
    y = data[target]
    x_train, x_test, y_train, y_test = train_test_split(x, y, test_size=0.25, random_state=44)
    tree = DecisionTreeRegressor(max_depth=4, min_samples_leaf=10, random_state=44)
    tree.fit(x_train, y_train)
    pred_train = tree.predict(x_train)
    pred_test = tree.predict(x_test)
    metrics = {
        "target": target,
        "rows": int(len(data)),
        "features": int(len(features)),
        "train_r2": float(r2_score(y_train, pred_train)),
        "test_r2": float(r2_score(y_test, pred_test)),
        "train_mae": float(mean_absolute_error(y_train, pred_train)),
        "test_mae": float(mean_absolute_error(y_test, pred_test)),
    }
    importance = pd.DataFrame({"feature": features, "importance": tree.feature_importances_})
    importance = importance.sort_values("importance", ascending=False).reset_index(drop=True)
    rules = export_text(tree, feature_names=features, max_depth=4)
    return tree, importance, metrics, rules


def make_figures(df: pd.DataFrame, ranges: list[dict], target_imp: pd.DataFrame, score_imp: pd.DataFrame, target_tree: DecisionTreeRegressor, features: list[str]) -> dict[str, Path]:
    figs: dict[str, Path] = {}

    fig, ax = plt.subplots(figsize=(8.0, 4.8))
    display = [
        ("spiral_turns", 0.75, 5.0, 1.0, 3.0, "turns"),
        ("eccentricity_e", 0.50, 1.0, 0.65, 1.0, "-"),
        ("channel_width", 50, 150, 80, 120, "um"),
        ("channel_height", 25, 100, 40, 75, "um"),
        ("inlet_velocity", 800, 5000, 1000, 3000, "um/s"),
        ("voltage", 4, 25, 6, 22, "V"),
        ("electrode_gap", 25, 100, 40, 80, "um"),
    ]
    for i, (name, lo, hi, pref_lo, pref_hi, unit) in enumerate(display):
        span = max(hi - lo, 1e-12)
        pref_a = (pref_lo - lo) / span
        pref_b = (pref_hi - lo) / span
        ax.plot([0, 1], [i, i], lw=10, color="#dbeafe", solid_capstyle="round")
        ax.plot([pref_a, pref_b], [i, i], lw=10, color=BLUE, solid_capstyle="round")
        ax.scatter([0, 1, pref_a, pref_b], [i, i, i, i], s=[18, 18, 26, 26], color=[MUTED, MUTED, BLUE, BLUE], zorder=4)
        ax.text(-0.02, i + 0.28, f"{lo:g}", ha="right", va="center", fontsize=7.8, color=MUTED)
        ax.text(1.02, i + 0.28, f"{hi:g} {unit}", ha="left", va="center", fontsize=7.8, color=MUTED)
        ax.text((pref_a + pref_b) / 2, i - 0.28, f"preferred {pref_lo:g}-{pref_hi:g}", ha="center", va="center", fontsize=7.5, color=INK)
    ax.set_yticks(range(len(display)), [d[0].replace("_", " ") for d in display])
    ax.invert_yaxis()
    ax.set_xlim(-0.15, 1.22)
    ax.set_xlabel("Normalized per-parameter range; light band = literature-bounded screen, dark band = preferred first pass")
    ax.set_title("Literature-bounded feasible design envelope")
    ax.spines[["top", "right", "left"]].set_visible(False)
    ax.grid(axis="x", alpha=0.25)
    figs["feasible_envelope"] = savefig(fig, "fig01_literature_feasible_envelope.png")

    fig, axes = plt.subplots(1, 2, figsize=(8.0, 4.2))
    for ax, imp, title, color in [
        (axes[0], target_imp.head(12), "Decision tree: target_correct", BLUE),
        (axes[1], score_imp.head(12), "Decision tree: design_score", GREEN),
    ]:
        top = imp.sort_values("importance")
        ax.barh(top["feature"], top["importance"], color=color)
        ax.set_xlabel("Gini-style regression importance")
        ax.set_title(title)
        ax.grid(axis="x", alpha=0.25)
    figs["tree_importance"] = savefig(fig, "fig02_decision_tree_importance.png")

    fig, ax = plt.subplots(figsize=(6.7, 4.8))
    sub = df.dropna(subset=["turns", "eccentricity", "target_correct"]).copy()
    sub["target_correct"] = pd.to_numeric(sub["target_correct"], errors="coerce")
    sc = ax.scatter(
        pd.to_numeric(sub["turns"], errors="coerce"),
        pd.to_numeric(sub["eccentricity"], errors="coerce"),
        c=sub["target_correct"],
        cmap="viridis",
        s=30,
        alpha=0.82,
        edgecolor="none",
    )
    fig.colorbar(sc, ax=ax, label="target_correct")
    ax.set_xlabel("Turns")
    ax.set_ylabel("Ellipticity e")
    ax.set_title("Observed simulation space: turn count versus ellipticity")
    ax.grid(alpha=0.22)
    figs["turn_ellipticity"] = savefig(fig, "fig03_turn_ellipticity_space.png")

    fig, axes = plt.subplots(1, 2, figsize=(8.0, 4.0))
    for ax, xcol, title in [
        (axes[0], "residence_s", "Accuracy versus residence time"),
        (axes[1], "effective_field_kV_m", "Accuracy versus effective field"),
    ]:
        if xcol not in df:
            ax.axis("off")
            continue
        valid = df.dropna(subset=[xcol, "target_correct"])
        ax.scatter(pd.to_numeric(valid[xcol], errors="coerce"), pd.to_numeric(valid["target_correct"], errors="coerce"), s=22, alpha=0.68, color=INK)
        ax.set_xlabel(xcol)
        ax.set_ylabel("target_correct")
        ax.set_title(title)
        ax.grid(alpha=0.25)
    fig.suptitle("Reality check: high accuracy can be driven by residence/field scaling", y=1.02)
    figs["reality_check"] = savefig(fig, "fig04_reality_check_accuracy_drivers.png")

    fig, ax = plt.subplots(figsize=(10.5, 5.2))
    plot_tree(
        target_tree,
        feature_names=features,
        filled=True,
        rounded=True,
        impurity=False,
        proportion=True,
        max_depth=3,
        fontsize=6,
        ax=ax,
    )
    ax.set_title("Interpretable decision tree for target_correct")
    figs["tree_structure"] = savefig(fig, "fig05_decision_tree_structure.png")
    return figs


def build_report(figs: dict[str, Path], ranges: list[dict], target_imp: pd.DataFrame, score_imp: pd.DataFrame, metrics: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        doc.styles[style_name].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Literature-Bounded Computational Design Envelope and Feature Importance for Spiral DEP Live/Dead Cell Sorting")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Feasibility and model-reality analysis for the next optimization stage")
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    def para(text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size: float = 10.5) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        rr = p.add_run(text)
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(size)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        rr = p.add_run(text.upper())
        rr.bold = True
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(11)

    def caption(text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rr = p.add_run(text)
        rr.italic = True
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(8.5)

    def picture(path: Path, cap: str, width: float = 15.8) -> None:
        doc.add_picture(str(path), width=Cm(width))
        caption(cap)

    heading("Abstract")
    para(
        "This report defines a literature-bounded design envelope and an interpretable feature-importance analysis for a continuous spiral dielectrophoretic device intended for live/dead cell sorting. The objective is not to claim final device performance, but to prevent the optimization from becoming an unconstrained numerical exercise. Feasible ranges are derived from spiral DEP, live/dead DEP, facing-electrode DEP, and thermal/ACET literature. Existing reduced-order simulation outputs are then combined and analyzed with decision-tree regressors to identify the dominant parameters controlling outlet classification and a penalized design score. The analysis indicates that channel width, velocity/residence-time variables, electrical field variables, inlet focusing, and geometry class currently dominate the reduced-order model. Turn count and ellipticity must therefore be isolated in the next experiment, rather than mixed with voltage and velocity changes."
    )

    heading("1. Literature-Bounded Feasible Envelope")
    para(
        "The feasible envelope was not selected arbitrarily. Betyar and Ramiar [R1] provide the closest spiral DEP/OpenFOAM precedent: one inlet, two outlets, side-wall electrodes, 100 um channel width, 1-5 spiral turns, and inlet velocities of 1000, 3000, and 5000 um/s. Elitas et al. [R5] provide a mammalian live/dead DEP anchor with 20 Vpp, 300 kHz, 1 uL/min, and approximately 90% dead-cell removal. Nguyen et al. [R10] provide a modern facing-electrode numerical workflow with explicit separation efficiency and purity, plus sweeps over voltage, velocity ratio, channel height, and electrode number. Jiang et al. [R15] motivate thermal gates because Joule heating and AC electrothermal flow can disturb DEP manipulation."
    )
    picture(figs["feasible_envelope"], "Figure 1. Literature-bounded feasible design envelope. The broad band is the screening range; the darker band is the preferred first-pass range for computational optimization.", 15.8)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Parameter", "Screen range", "Preferred range", "Refs", "Reason"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in ranges:
        cells = table.add_row().cells
        cells[0].text = f"{row['parameter']} ({row['units']})"
        cells[1].text = row["proposed_screen_range"]
        cells[2].text = row["preferred_initial_range"]
        cells[3].text = row["primary_refs"]
        cells[4].text = row["design_reason"]
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for rr in p.runs:
                    rr.font.name = "Times New Roman"
                    rr.font.size = Pt(7.8)
    caption("Table 1. Feasible parameter ranges extracted from the local literature set and converted into optimization bounds.")

    heading("2. Mathematical Model Stored in LaTeX")
    para(
        "All equations used by the computational model are stored in results/final_v1/latex_equations.tex. The next stage should keep the formula family deliberately simple: an Archimedean spiral with elliptic scaling. This makes turn count and ellipticity testable instead of hiding them inside arbitrary curves."
    )
    for eq in [
        r"r(\theta)=r_0+p\theta,\quad 0\leq\theta\leq2\pi N",
        r"x(\theta)=r(\theta)\cos\theta,\quad y(\theta)=e\,r(\theta)\sin\theta",
        r"\mathbf{F}_{DEP}=2\pi\epsilon_m r_p^3\operatorname{Re}\{K(\omega)\}\nabla|\mathbf{E}|^2",
        r"K(\omega)=\frac{\epsilon_p^\ast-\epsilon_m^\ast}{\epsilon_p^\ast+2\epsilon_m^\ast},\quad \epsilon^\ast=\epsilon-\frac{j\sigma}{\omega}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(eq)
        rr.font.name = "Courier New"
        rr.font.size = Pt(9.5)

    heading("3. Feature Catalogue and Decision-Tree Method")
    para(
        "The feature catalogue separates geometry, electrical, flow, inlet/outlet, and output variables. A decision-tree regressor was fitted to existing screening simulations from the clean-shape and final operating datasets. Two targets were analyzed: raw target_correct and a penalized design_score that subtracts wall-loss, power, length, and thermal-risk penalties. This is intentionally interpretable; the purpose is to expose model dependence, not to produce a black-box optimizer."
    )
    para(
        f"The decision tree for target_correct used {metrics['target_correct']['rows']} rows and {metrics['target_correct']['features']} numeric features. Test R2 was {metrics['target_correct']['test_r2']:.3f} with test MAE {metrics['target_correct']['test_mae']:.3f}. Because this is observational reduced-order data, feature importance should be read as model sensitivity, not causal proof."
    )
    picture(figs["tree_importance"], "Figure 2. Decision-tree feature importance for raw target_correct and penalized design_score.", 15.8)
    picture(figs["tree_structure"], "Figure 3. Interpretable decision-tree structure for target_correct. The tree is intentionally shallow to avoid opaque overfitting.", 15.8)

    heading("4. Turn Count and Ellipticity Reality Check")
    para(
        "The current data do include variation in turn count and ellipticity, but those variables are not yet isolated cleanly. They are mixed with geometry family, velocity, width, voltage, active DEP coverage, and residence time. Therefore, a high-accuracy candidate cannot yet be used as evidence that a larger number of turns or a more elliptical spiral is intrinsically superior."
    )
    picture(figs["turn_ellipticity"], "Figure 4. Observed simulation space for turn count and ellipticity. This plot motivates a controlled turn-by-ellipticity sweep.", 14.5)
    picture(figs["reality_check"], "Figure 5. Reality check showing why accuracy alone is insufficient: residence time and field strength can drive apparent separation in the reduced-order model.", 15.8)

    heading("5. Most Important Features")
    top = target_imp.head(8)
    for _, row in top.iterrows():
        para(f"- {row['feature']}: importance {row['importance']:.3f}", WD_ALIGN_PARAGRAPH.LEFT, 10)
    para(
        "The immediate implication is that the next experiment must control confounding. Turn count should be swept at fixed voltage, velocity, width, frequency, electrode gap, and outlet split; ellipticity should be swept at fixed length or analyzed with length-normalized performance. Otherwise, the optimizer will rediscover the trivial rule: increase residence time or effective field until particles separate."
    )

    heading("6. Proposed Next Computational Experiment")
    para(
        "The next design experiment should be a structured grid rather than a broad random search. Recommended variables are: turns N in {0.75, 1.0, 1.5, 2.0, 2.5, 3.0, 4.0, 5.0}; ellipticity e in {0.55, 0.70, 0.85, 1.00}; voltage in a narrow feasible band around the literature-supported low-voltage region; and velocity fixed initially at 1000 and 3000 um/s. Each point should be compared with same-length straight DEP and no-DEP controls. Only after the turn/eccentricity mechanism is clean should ML be used for fine tuning."
    )

    heading("7. References")
    refs = [
        "[R1] S. Betyar and A. Ramiar, Numerical Simulation of Biological Particle Separation in a Spiral Microchannel Using the Dielectrophoresis Mechanism, 2026.",
        "[R2] N. Lewpiriyawong, C. Yang, and X. Xuan, Particle separation by charge in spiral microchannels, Biomicrofluidics, 2011.",
        "[R5] Y. Yildizhan et al., Dielectrophoretic Separation of Live and Dead Monocytes Using 3D Carbon-Electrodes, Sensors, 2017.",
        "[R9] Huang et al., Label-free Live and Dead Cell Separation Method Using a High-Efficiency ODEP Force-based Microfluidic Platform, 2025.",
        "[R10] Nguyen et al., Facing-electrode DEP cell separation, Scientific Reports, 2024.",
        "[R15] Jiang et al., Numerical simulations of combined DEP and AC electrothermal flow, Micromachines, 2024.",
        "[R16] Petruzzellis et al., Lab-on-Chip Systems for Cell Sorting and Spiral Microchannels, Micromachines, 2024.",
    ]
    for ref in refs:
        para(ref, WD_ALIGN_PARAGRAPH.LEFT, 9.2)

    doc.save(REPORT_DOCX)


def render_docx() -> tuple[Path | None, list[Path]]:
    render_dir = OUT_DIR / "docx_render"
    render_dir.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation=file:///tmp/lo_profile_final_v1_{REPORT_DOCX.stem}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(render_dir),
                str(REPORT_DOCX),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        pdf = render_dir / f"{REPORT_DOCX.stem}.pdf"
        if pdf.exists():
            subprocess.run(["pdftoppm", "-png", str(pdf), str(render_dir / "page")], check=True)
            return pdf, sorted(render_dir.glob("page-*.png"))
    except (OSError, subprocess.CalledProcessError):
        return None, []
    return None, []


def write_readme(figs: dict[str, Path], metrics: dict, pdf: Path | None, pages: list[Path]) -> None:
    lines = [
        "# Final Feasibility and Feature Analysis",
        "",
        "This package defines a literature-bounded feasible envelope and a decision-tree feature-importance reality check for the next spiral DEP live/dead sorting optimization.",
        "",
        "## Key conclusion",
        "",
        "The current reduced-order data are useful for screening, but turn count and ellipticity must be isolated in a controlled experiment before claiming spiral-specific mechanism. Decision trees show that width, velocity/residence, field variables, inlet focusing, and geometry class dominate the current model.",
        "",
        "## Outputs",
        "",
        f"- DOCX report: `{REPORT_DOCX.relative_to(ROOT)}`",
        f"- rendered PDF: `{pdf.relative_to(ROOT) if pdf else 'not rendered'}`",
        f"- rendered page count: `{len(pages)}`",
        "- feasible ranges: `results/final_v1/literature_feasible_ranges.csv`",
        "- equations: `results/final_v1/latex_equations.tex`",
        "- feature catalogue: `results/final_v1/feature_catalog.csv`",
        "- feature importance: `results/final_v1/decision_tree_feature_importance.csv`",
        "- tree rules: `results/final_v1/decision_tree_rules_target_correct.txt`, `results/final_v1/decision_tree_rules_design_score.txt`",
        "- model diagnostics: `results/final_v1/model_diagnostics.json`",
        "",
        "## Diagnostics",
        "",
        f"- target_correct decision tree test R2: `{metrics['target_correct']['test_r2']:.3f}`",
        f"- target_correct decision tree test MAE: `{metrics['target_correct']['test_mae']:.3f}`",
        f"- design_score decision tree test R2: `{metrics['design_score']['test_r2']:.3f}`",
        f"- design_score decision tree test MAE: `{metrics['design_score']['test_mae']:.3f}`",
        "",
        "## Figures",
    ]
    for key, path in figs.items():
        lines.append(f"- `{key}`: `{path.relative_to(ROOT)}`")
    (OUT_DIR / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    ensure_dirs()
    set_style()

    ranges = feasible_ranges()
    write_rows(OUT_DIR / "literature_feasible_ranges.csv", ranges)
    (OUT_DIR / "latex_equations.tex").write_text(latex_equations(), encoding="utf-8")
    write_rows(OUT_DIR / "feature_catalog.csv", feature_catalog())

    df = load_ml_dataset()
    df.to_csv(OUT_DIR / "combined_screening_dataset.csv", index=False)

    target_tree, target_imp, target_metrics, target_rules = train_decision_tree(df, "target_correct")
    score_tree, score_imp, score_metrics, score_rules = train_decision_tree(df, "design_score")
    target_imp["target"] = "target_correct"
    score_imp["target"] = "design_score"
    importance = pd.concat([target_imp, score_imp], ignore_index=True)
    importance.to_csv(OUT_DIR / "decision_tree_feature_importance.csv", index=False)
    (OUT_DIR / "decision_tree_rules_target_correct.txt").write_text(target_rules, encoding="utf-8")
    (OUT_DIR / "decision_tree_rules_design_score.txt").write_text(score_rules, encoding="utf-8")
    metrics = {"target_correct": target_metrics, "design_score": score_metrics}
    (OUT_DIR / "model_diagnostics.json").write_text(json.dumps(metrics, indent=2), encoding="utf-8")

    features = numeric_features(df)
    figs = make_figures(df, ranges, target_imp, score_imp, target_tree, features)
    build_report(figs, ranges, target_imp, score_imp, metrics)
    pdf, pages = render_docx()
    write_readme(figs, metrics, pdf, pages)
    print(json.dumps({"out_dir": str(OUT_DIR), "report_docx": str(REPORT_DOCX), "rendered_pdf": str(pdf) if pdf else None, "pages": len(pages)}, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
