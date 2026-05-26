#!/usr/bin/env python3
"""V14 IEEE-style publication package with device-level figures.

V14 does not invent new numerical performance. It repackages the V12/V13
validated reduced-order results with publication-grade device figures:
channel ribbon, side electrodes, inlet, outlet splitter, dimensions, and
representative particle trajectories. The geometry library panel is a visual
design-space record; only the V12 selected design carries validated metrics.
"""

from __future__ import annotations

import csv
import json
import math
import shutil
import subprocess
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from matplotlib.patches import FancyArrowPatch, Polygon, Rectangle

from design_v11_shape_ml_nonintersection import centerline, path_length, sample_to_spec


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "results" / "design_v14_ieee_publication_package"
FIG_DIR = OUT_DIR / "figures"
DOC_DIR = ROOT / "output" / "doc"
REPORT_DOCX = DOC_DIR / "spiral_dep_v14_ieee_publication_draft.docx"


BLUE = "#2563eb"
GREEN = "#16a34a"
RED = "#dc2626"
INK = "#0f172a"
MUTED = "#64748b"
CHANNEL = "#eef2f7"
EDGE = "#1e293b"
GOLD = "#f59e0b"


def read_csv(path: Path) -> list[dict]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def write_csv(path: Path, rows: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()), lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


def f(row: dict, key: str, default: float = 0.0) -> float:
    try:
        return float(row.get(key, default))
    except (TypeError, ValueError):
        return default


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def final_sample_from_v12() -> tuple[dict, dict]:
    sample = load_json(ROOT / "results" / "design_v11_shape_ml_nonintersection" / "rank1_monotone_curvature_spiral_formula.json")
    final = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_final_validation_summary.csv")[0]
    sample["voltage_v"] = f(final, "voltage_v")
    sample["velocity_m_s"] = f(final, "velocity_um_s") * 1e-6
    sample["electrode_gap_m"] = f(final, "electrode_gap_um") * 1e-6
    sample["electrode_coverage"] = f(final, "electrode_coverage")
    sample["frequency_hz"] = f(final, "frequency_hz")
    return sample, final


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DOC_DIR.mkdir(parents=True, exist_ok=True)


def set_pub_style() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 8.5,
            "axes.titlesize": 9.5,
            "axes.labelsize": 8.5,
            "legend.fontsize": 7.5,
            "xtick.labelsize": 7.5,
            "ytick.labelsize": 7.5,
            "axes.linewidth": 0.8,
            "savefig.dpi": 420,
        }
    )


def savefig(fig: plt.Figure, name: str, tight: bool = True) -> Path:
    path = FIG_DIR / name
    if tight:
        fig.tight_layout()
    fig.savefig(path, dpi=420, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def normals(x: np.ndarray, y: np.ndarray) -> tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
    dx = np.gradient(x)
    dy = np.gradient(y)
    mag = np.maximum(np.hypot(dx, dy), 1e-15)
    tx, ty = dx / mag, dy / mag
    nx, ny = -ty, tx
    return tx, ty, nx, ny


def ribbon_bounds(x: np.ndarray, y: np.ndarray, width: float) -> tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
    _, _, nx, ny = normals(x, y)
    left_x = x + 0.5 * width * nx
    left_y = y + 0.5 * width * ny
    right_x = x - 0.5 * width * nx
    right_y = y - 0.5 * width * ny
    return left_x, left_y, right_x, right_y


def draw_channel(
    ax: plt.Axes,
    x: np.ndarray,
    y: np.ndarray,
    width: float,
    active: tuple[float, float] = (0.22, 0.80),
    show_center: bool = False,
    electrode_lw: float = 3.0,
    alpha: float = 1.0,
) -> dict[str, np.ndarray]:
    lx, ly, rx, ry = ribbon_bounds(x, y, width)
    polygon = np.column_stack([np.r_[lx, rx[::-1]], np.r_[ly, ry[::-1]]])
    ax.add_patch(Polygon(polygon * 1e3, closed=True, facecolor=CHANNEL, edgecolor="none", alpha=alpha, zorder=1))
    ax.plot(lx * 1e3, ly * 1e3, color=EDGE, lw=1.25, zorder=3)
    ax.plot(rx * 1e3, ry * 1e3, color=EDGE, lw=1.25, zorder=3)
    if show_center:
        ax.plot(x * 1e3, y * 1e3, color=MUTED, lw=0.5, ls=":", zorder=2)

    n = len(x)
    i0 = int(np.clip(active[0], 0.0, 1.0) * (n - 1))
    i1 = int(np.clip(active[1], 0.0, 1.0) * (n - 1))
    if i1 > i0:
        ax.plot(lx[i0:i1] * 1e3, ly[i0:i1] * 1e3, color=BLUE, lw=electrode_lw, solid_capstyle="round", zorder=5)
        ax.plot(rx[i0:i1] * 1e3, ry[i0:i1] * 1e3, color=GOLD, lw=electrode_lw, solid_capstyle="round", zorder=5)
    return {"lx": lx, "ly": ly, "rx": rx, "ry": ry}


def add_inlet_outlet(ax: plt.Axes, x: np.ndarray, y: np.ndarray, width: float, labels: bool = True) -> None:
    tx, ty, nx, ny = normals(x, y)
    start = np.array([x[0], y[0]])
    end = np.array([x[-1], y[-1]])
    inlet_from = start - np.array([tx[0], ty[0]]) * 0.55e-3
    inlet_to = start + np.array([tx[0], ty[0]]) * 0.08e-3
    ax.add_patch(
        FancyArrowPatch(
            inlet_from * 1e3,
            inlet_to * 1e3,
            arrowstyle="-|>",
            mutation_scale=12,
            lw=1.8,
            color=INK,
            zorder=8,
        )
    )
    out_base = end + np.array([tx[-1], ty[-1]]) * 0.08e-3
    out_live = end + np.array([tx[-1], ty[-1]]) * 0.58e-3 + np.array([nx[-1], ny[-1]]) * width * 0.85
    out_dead = end + np.array([tx[-1], ty[-1]]) * 0.58e-3 - np.array([nx[-1], ny[-1]]) * width * 0.85
    ax.add_patch(FancyArrowPatch(out_base * 1e3, out_live * 1e3, arrowstyle="-|>", mutation_scale=11, lw=1.7, color=GREEN, zorder=8))
    ax.add_patch(FancyArrowPatch(out_base * 1e3, out_dead * 1e3, arrowstyle="-|>", mutation_scale=11, lw=1.7, color=RED, zorder=8))
    if labels:
        label_live = out_live + np.array([tx[-1], ty[-1]]) * 0.12e-3 + np.array([nx[-1], ny[-1]]) * 0.16e-3
        label_dead = out_dead + np.array([tx[-1], ty[-1]]) * 0.12e-3 - np.array([nx[-1], ny[-1]]) * 0.16e-3
        label_box = dict(boxstyle="round,pad=0.14", facecolor="white", edgecolor="none", alpha=0.84)
        ax.text(*(inlet_from * 1e3), "Inlet", ha="right", va="center", color=INK, fontsize=8.5, bbox=label_box)
        ax.text(*(label_live * 1e3), "Outlet A\nlive-rich", ha="left", va="center", color=GREEN, fontsize=7.3, bbox=label_box)
        ax.text(*(label_dead * 1e3), "Outlet B\ndead-rich", ha="left", va="center", color=RED, fontsize=7.3, bbox=label_box)


def format_device_axis(ax: plt.Axes) -> None:
    ax.set_aspect("equal")
    ax.set_xlabel("x (mm)")
    ax.set_ylabel("y (mm)")
    ax.grid(False)
    ax.spines[["top", "right"]].set_visible(False)


def metric_box(ax: plt.Axes, text: str, loc: tuple[float, float] = (0.02, 0.02)) -> None:
    ax.text(
        loc[0],
        loc[1],
        text,
        transform=ax.transAxes,
        ha="left",
        va="bottom",
        fontsize=7.4,
        color=INK,
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white", edgecolor="#cbd5e1", alpha=0.92),
        zorder=20,
    )


def mapped_paths(sample: dict, rows: list[dict]) -> tuple[np.ndarray, np.ndarray]:
    _, xc, yc = centerline(sample, n=2200)
    _, _, nx, ny = normals(xc, yc)
    xs, ys = [], []
    for row in rows:
        p = float(row["progress"])
        idx = min(len(xc) - 1, max(0, int(round(p * (len(xc) - 1)))))
        lateral = float(row["lateral_m"])
        xs.append(xc[idx] + nx[idx] * lateral)
        ys.append(yc[idx] + ny[idx] * lateral)
    return np.array(xs), np.array(ys)


def geometry_library_centerline(kind: str, n: int = 720) -> tuple[np.ndarray, np.ndarray, float, str]:
    theta = np.linspace(0.0, 2.2 * math.pi, n)
    width = 35e-6
    formula = ""
    if kind == "archimedean":
        r = 0.35e-3 + 0.092e-3 * theta
        x, y = r * np.cos(theta), r * np.sin(theta)
        formula = "r=a+b theta"
    elif kind == "elliptic_archimedean":
        r = 0.38e-3 + 0.088e-3 * theta
        x, y = r * np.cos(theta), 0.66 * r * np.sin(theta)
        formula = "x=r cos(theta), y=e r sin(theta)"
    elif kind == "monotone_curvature":
        theta = np.linspace(0.0, 1.85 * math.pi, n)
        r = 1.12e-3 / (1.0 + 0.105 * theta) ** 1.08
        x, y = r * np.cos(theta), r * np.sin(theta)
        formula = "r=r0/(1+k theta)^p"
    elif kind == "logarithmic":
        theta = np.linspace(0.0, 1.85 * math.pi, n)
        r = 0.34e-3 * np.exp(0.145 * theta)
        x, y = r * np.cos(theta), r * np.sin(theta)
        formula = "r=a exp(k theta)"
    elif kind == "fermat":
        r = 0.38e-3 * np.sqrt(1.0 + 0.72 * theta)
        x, y = r * np.cos(theta), r * np.sin(theta)
        formula = "r=a sqrt(1+k theta)"
    elif kind == "superellipse_oval":
        r = 0.38e-3 + 0.082e-3 * theta
        c, s = np.cos(theta), np.sin(theta)
        q = 0.72
        x = r * np.sign(c) * np.abs(c) ** q
        y = 0.70 * r * np.sign(s) * np.abs(s) ** q
        formula = "superellipse spiral"
    elif kind == "racetrack_oval":
        r = 0.42e-3 + 0.080e-3 * theta
        x = r * np.cos(theta) + 0.22e-3 * np.tanh(2.2 * np.cos(theta))
        y = 0.68 * r * np.sin(theta)
        formula = "oval/stadium mapping"
    elif kind == "short_c_arc":
        theta = np.linspace(-0.15 * math.pi, 1.18 * math.pi, n)
        r = 0.68e-3 + 0.035e-3 * theta
        x, y = r * np.cos(theta), r * np.sin(theta)
        formula = "short C-arc"
    elif kind == "elliptic_c_arc":
        theta = np.linspace(-0.20 * math.pi, 1.28 * math.pi, n)
        r = 0.72e-3 + 0.045e-3 * theta
        x, y = r * np.cos(theta), 0.62 * r * np.sin(theta)
        formula = "elliptic C-arc"
    elif kind == "s_bend_spiral":
        theta2 = np.linspace(0.0, 1.42 * math.pi, n - 130)
        r = 0.42e-3 + 0.090e-3 * theta2
        xs, ys = r * np.cos(theta2), 0.72 * r * np.sin(theta2)
        u = np.linspace(0.0, 1.0, 130)
        inlet_x = xs[0] - 0.55e-3 * (1.0 - u)
        inlet_y = ys[0] + 0.18e-3 * np.sin(math.pi * u)
        x, y = np.r_[inlet_x[:-1], xs], np.r_[inlet_y[:-1], ys]
        formula = "S-bend prefocus + spiral"
    elif kind == "contraction_spiral":
        theta2 = np.linspace(0.0, 1.65 * math.pi, n - 100)
        r = 0.42e-3 + 0.085e-3 * theta2
        xs, ys = r * np.cos(theta2), 0.74 * r * np.sin(theta2)
        inlet_x = np.linspace(xs[0] - 0.62e-3, xs[0], 100)
        inlet_y = np.linspace(ys[0] - 0.08e-3, ys[0], 100)
        x, y = np.r_[inlet_x[:-1], xs], np.r_[inlet_y[:-1], ys]
        formula = "contraction inlet + elliptic spiral"
    elif kind == "linked_oval_arc":
        theta = np.linspace(-0.15 * math.pi, 1.72 * math.pi, n)
        r = 0.48e-3 + 0.040e-3 * theta
        x = r * np.cos(theta) + 0.16e-3 * np.sin(0.5 * theta)
        y = 0.58 * r * np.sin(theta)
        formula = "linked oval arc, no crossover"
    elif kind == "spiral_expanded_outlet":
        r = 0.36e-3 + 0.090e-3 * theta
        x, y = r * np.cos(theta), 0.70 * r * np.sin(theta)
        y[-90:] += np.linspace(0.0, 0.20e-3, 90)
        formula = "elliptic spiral + expanded outlet"
    elif kind == "kidney_spiral":
        r = 0.40e-3 + 0.084e-3 * theta
        x = r * np.cos(theta) * (1.0 + 0.08 * np.sin(theta))
        y = 0.70 * r * np.sin(theta) * (1.0 - 0.05 * np.cos(theta))
        formula = "smooth kidney-like oval"
    elif kind == "D_spiral":
        r = 0.40e-3 + 0.082e-3 * theta
        x = r * np.cos(theta)
        y = 0.64 * r * np.sin(theta)
        x = x + 0.16e-3 * (1.0 / (1.0 + np.exp(-6.0 * np.cos(theta))) - 0.5)
        formula = "D-shaped smooth spiral"
    elif kind == "teardrop_spiral":
        theta = np.linspace(0.0, 1.86 * math.pi, n)
        r = 0.42e-3 + 0.080e-3 * theta
        taper = 1.0 + 0.18 * np.sin(theta - 0.45)
        x = r * np.cos(theta) * taper
        y = 0.68 * r * np.sin(theta) * (1.0 - 0.10 * np.cos(theta))
        formula = "smooth teardrop spiral"
    elif kind == "spiral_with_arc_inlet":
        theta2 = np.linspace(0.0, 1.55 * math.pi, n - 120)
        r = 0.44e-3 + 0.083e-3 * theta2
        xs, ys = r * np.cos(theta2), 0.68 * r * np.sin(theta2)
        a = np.linspace(-0.9, 0.0, 120)
        inlet_x = xs[0] - 0.48e-3 * np.cos(a)
        inlet_y = ys[0] + 0.20e-3 * np.sin(a)
        x, y = np.r_[inlet_x[:-1], xs], np.r_[inlet_y[:-1], ys]
        formula = "arc inlet + elliptic spiral"
    else:
        raise ValueError(kind)
    return x, y, width, formula


def min_clearance_ratio(x: np.ndarray, y: np.ndarray, width: float) -> float:
    pts = np.column_stack([x, y])
    skip = max(10, len(pts) // 30)
    best = float("inf")
    for i in range(0, len(pts), 3):
        j0 = min(len(pts), i + skip)
        if j0 >= len(pts):
            continue
        d = np.hypot(pts[j0:, 0] - pts[i, 0], pts[j0:, 1] - pts[i, 1])
        if len(d):
            best = min(best, float(np.min(d)))
    return best / width if math.isfinite(best) else 99.0


def make_final_device_figure(sample: dict, final: dict) -> Path:
    _, x, y = centerline(sample, n=2600)
    width = sample["channel_width_m"]
    active = (sample["dep_start_fraction"], sample["dep_end_fraction"])
    fig = plt.figure(figsize=(7.25, 4.55))
    gs = fig.add_gridspec(2, 3, width_ratios=[2.2, 0.9, 0.9], height_ratios=[1, 1])
    ax = fig.add_subplot(gs[:, 0])
    draw_channel(ax, x, y, width, active, electrode_lw=4.2)
    add_inlet_outlet(ax, x, y, width, labels=False)
    ax.set_title("Selected compact spiral DEP sorter")
    format_device_axis(ax)
    metric_box(
        ax,
        "\n".join(
            [
                f"L = {path_length(sample)*1e3:.2f} mm",
                f"W = {width*1e6:.0f} um, gap = {sample['electrode_gap_m']*1e6:.1f} um",
                f"V = {sample['voltage_v']:.1f} V, f = {sample['frequency_hz']/1e6:.2f} MHz",
                f"Q = {f(final, 'flow_uL_min'):.3f} uL/min",
            ]
        ),
        (0.03, 0.04),
    )
    ax.plot([], [], color=BLUE, lw=3, label="Electrode 1")
    ax.plot([], [], color=GOLD, lw=3, label="Electrode 2")
    ax.legend(loc="upper left", frameon=True, framealpha=0.95)

    for sub_ax, frac, title in [
        (fig.add_subplot(gs[0, 1]), 0.03, "Inlet prefocus"),
        (fig.add_subplot(gs[0, 2]), 0.50, "Active DEP zone"),
        (fig.add_subplot(gs[1, 1]), 0.82, "Outlet split"),
    ]:
        n = len(x)
        i = int(frac * (n - 1))
        lo, hi = max(0, i - 65), min(n, i + 65)
        draw_channel(sub_ax, x[lo:hi], y[lo:hi], width, (0.0, 1.0), electrode_lw=3.0)
        if title == "Outlet split":
            add_inlet_outlet(sub_ax, x[lo:hi], y[lo:hi], width, labels=False)
        sub_ax.set_aspect("equal")
        sub_ax.axis("off")
        sub_ax.set_title(title, pad=2)

    ax_dim = fig.add_subplot(gs[1, 2])
    ax_dim.axis("off")
    ax_dim.add_patch(Rectangle((0.12, 0.46), 0.70, 0.18, facecolor=CHANNEL, edgecolor=EDGE, lw=1.0))
    ax_dim.plot([0.12, 0.82], [0.71, 0.71], color=BLUE, lw=4, solid_capstyle="round")
    ax_dim.plot([0.12, 0.82], [0.39, 0.39], color=GOLD, lw=4, solid_capstyle="round")
    ax_dim.annotate("", xy=(0.10, 0.46), xytext=(0.10, 0.64), arrowprops=dict(arrowstyle="<->", lw=1.0, color=INK))
    ax_dim.text(0.03, 0.55, "W", ha="center", va="center", fontsize=8)
    ax_dim.annotate("", xy=(0.86, 0.39), xytext=(0.86, 0.71), arrowprops=dict(arrowstyle="<->", lw=1.0, color=INK))
    ax_dim.text(0.93, 0.55, "g", ha="center", va="center", fontsize=8)
    ax_dim.text(0.47, 0.16, "side-wall electrode abstraction", ha="center", va="center", fontsize=7.5, color=MUTED)
    return savefig(fig, "fig01_final_device_schematic.png")


def make_geometry_library_panel() -> tuple[Path, Path]:
    kinds = [
        "archimedean",
        "elliptic_archimedean",
        "monotone_curvature",
        "logarithmic",
        "fermat",
        "superellipse_oval",
        "racetrack_oval",
        "short_c_arc",
        "elliptic_c_arc",
        "s_bend_spiral",
        "contraction_spiral",
        "linked_oval_arc",
        "spiral_expanded_outlet",
        "kidney_spiral",
        "D_spiral",
        "teardrop_spiral",
    ]
    rows: list[dict] = []
    fig, axes = plt.subplots(4, 4, figsize=(8.1, 7.9))
    for ax, kind in zip(axes.flat, kinds):
        x, y, width, formula = geometry_library_centerline(kind)
        clearance = min_clearance_ratio(x, y, width)
        length = float(np.sum(np.hypot(np.diff(x), np.diff(y)))) * 1e3
        bbox = (x.max() - x.min() + width) * (y.max() - y.min() + width) * 1e6
        valid = clearance > 1.00
        draw_channel(ax, x, y, width, (0.18, 0.78), electrode_lw=2.2)
        add_inlet_outlet(ax, x, y, width, labels=False)
        ax.set_aspect("equal")
        ax.axis("off")
        ax.set_title(f"{kind.replace('_', ' ')}\nL={length:.1f} mm, clr={clearance:.1f}W", fontsize=7.2)
        rows.append(
            {
                "geometry": kind,
                "formula": formula,
                "length_mm": f"{length:.4f}",
                "footprint_mm2": f"{bbox:.4f}",
                "min_clearance_widths": f"{clearance:.4f}",
                "passes_visual_clearance_gate": str(valid),
                "purpose": "publication geometry library; not final simulated metric unless selected",
            }
        )
    fig.suptitle("V14 formula-defined device geometry library with channel ribbons and side electrodes", y=0.995, fontsize=10.5)
    panel_path = savefig(fig, "fig02_geometry_library_panel.png", tight=True)
    csv_path = OUT_DIR / "v14_geometry_library.csv"
    write_csv(csv_path, rows)
    return panel_path, csv_path


def make_trajectory_device_overlay(sample: dict, final: dict) -> Path:
    path_file = ROOT / "results" / "design_v13_manuscript_report" / "v13_final_trajectory_paths.csv"
    rows = read_csv(path_file)
    _, x, y = centerline(sample, n=2600)
    width = sample["channel_width_m"]
    fig, ax = plt.subplots(figsize=(7.15, 4.95))
    draw_channel(ax, x, y, width, (sample["dep_start_fraction"], sample["dep_end_fraction"]), electrode_lw=3.4)
    add_inlet_outlet(ax, x, y, width, labels=False)
    for cls, color, step in [("live", GREEN, 3), ("dead", RED, 3)]:
        ids = sorted({int(r["particle_id"]) for r in rows if r["class"] == cls})
        for pid in ids[::step][:10]:
            pr = [r for r in rows if r["class"] == cls and int(r["particle_id"]) == pid]
            tx, ty = mapped_paths(sample, pr)
            ax.plot(tx * 1e3, ty * 1e3, color=color, lw=0.95, alpha=0.56, zorder=7)
    ax.plot([], [], color=GREEN, lw=1.4, label="live trajectories")
    ax.plot([], [], color=RED, lw=1.4, label="dead trajectories")
    ax.set_title("Representative simulated trajectories on the selected device")
    format_device_axis(ax)
    ax.legend(loc="upper right", frameon=True, framealpha=0.95)
    metric_box(
        ax,
        f"Target correct = {f(final, 'mean_target_correct'):.3f} +/- {f(final, 'std_target_correct'):.3f}\n"
        f"Topology gain = {f(final, 'topology_gain_vs_straight'):.3f}\n"
        f"Wall loss = {f(final, 'mean_wall_loss'):.3f}",
        (0.03, 0.04),
    )
    return savefig(fig, "fig03_device_trajectory_overlay.png")


def make_results_plate(final: dict) -> Path:
    v12 = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_final_validation_summary.csv")
    screen = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_operating_screen.csv")
    fig = plt.figure(figsize=(7.25, 5.15))
    gs = fig.add_gridspec(2, 2, hspace=0.36, wspace=0.34)

    ax = fig.add_subplot(gs[0, 0])
    metrics = ["mean_target_correct", "mean_live_recovery", "mean_dead_removal", "mean_live_outlet_purity", "mean_dead_outlet_purity"]
    labels = ["Correct", "Live\nrecovery", "Dead\nremoval", "Live\npurity", "Dead\npurity"]
    vals = [f(final, m) for m in metrics]
    ax.bar(labels, vals, color=[INK, GREEN, RED, "#0ea5e9", "#f97316"], width=0.72)
    ax.axhline(0.80, color=MUTED, ls="--", lw=0.9)
    ax.set_ylim(0, 1.08)
    ax.set_ylabel("Metric")
    ax.set_title("Final validation metrics")

    ax = fig.add_subplot(gs[0, 1])
    ax.scatter([f(r, "voltage_v") for r in screen], [f(r, "target_correct") for r in screen], c=[f(r, "active_joule_power_mW") for r in screen], cmap="magma", s=18, alpha=0.72, edgecolor="none")
    ax.scatter([f(final, "voltage_v")], [f(final, "mean_target_correct")], s=84, color="#22c55e", edgecolor=INK, lw=0.8, zorder=5)
    ax.set_xlabel("Voltage (V)")
    ax.set_ylabel("Correct")
    ax.set_title("Operating screen")

    ax = fig.add_subplot(gs[1, 0])
    controls = ["no_dep_correct", "unfocused_inlet_correct", "straight_control_correct", "mean_target_correct"]
    control_labels = ["No DEP", "Unfocused", "Straight", "Selected"]
    ax.plot(control_labels, [f(final, c) for c in controls], marker="o", color=INK, lw=1.7)
    ax.set_ylim(0.45, 1.02)
    ax.set_ylabel("Correct")
    ax.set_title("Mechanism controls")
    ax.tick_params(axis="x", rotation=18)

    ax = fig.add_subplot(gs[1, 1])
    thermo = ["active_joule_power_mW", "steady_substrate_delta_c_proxy", "pressure_drop_kPa", "flow_uL_min"]
    tlabels = ["Power\n(mW)", "Delta T\n(C)", "dP\n(kPa)", "Flow\nuL/min"]
    ax.bar(tlabels, [f(final, k) for k in thermo], color=["#f97316", "#ef4444", "#6366f1", "#14b8a6"])
    ax.set_title("Device-operation proxies")
    ax.set_ylabel("Value")
    return savefig(fig, "fig04_ieee_results_plate.png")


def make_outlet_distribution() -> Path:
    particle_file = ROOT / "results" / "design_v13_manuscript_report" / "v13_final_trajectory_particles.csv"
    rows = read_csv(particle_file)
    fig, ax = plt.subplots(figsize=(3.55, 2.65))
    live = [f(r, "final_lateral_um") for r in rows if r["class"] == "live"]
    dead = [f(r, "final_lateral_um") for r in rows if r["class"] == "dead"]
    bins = np.linspace(min(live + dead) - 1.0, max(live + dead) + 1.0, 24)
    ax.hist(live, bins=bins, alpha=0.70, color=GREEN, label="live")
    ax.hist(dead, bins=bins, alpha=0.70, color=RED, label="dead")
    ax.set_xlabel("Final lateral position (um)")
    ax.set_ylabel("Count")
    ax.set_title("Outlet-position distribution")
    ax.legend(frameon=False)
    return savefig(fig, "fig05_outlet_distribution.png")


def copy_reference_plots() -> dict[str, Path]:
    sources = {
        "workflow": ROOT / "results" / "design_v12_final_package" / "figures" / "fig12_workflow.png",
        "pareto": ROOT / "results" / "design_v12_final_package" / "figures" / "fig05_operating_pareto_power.png",
        "thermal": ROOT / "results" / "design_v12_final_package" / "figures" / "fig07_thermal_map.png",
        "feature_importance": ROOT / "results" / "design_v12_final_package" / "figures" / "fig04_ml_feature_importance.png",
    }
    out: dict[str, Path] = {}
    for key, src in sources.items():
        dst = FIG_DIR / f"ref_{key}.png"
        shutil.copy2(src, dst)
        out[key] = dst
    return out


def build_figures() -> dict[str, Path]:
    ensure_dirs()
    set_pub_style()
    sample, final = final_sample_from_v12()
    figs = copy_reference_plots()
    figs["device"] = make_final_device_figure(sample, final)
    figs["library"], _ = make_geometry_library_panel()
    figs["trajectory"] = make_trajectory_device_overlay(sample, final)
    figs["results_plate"] = make_results_plate(final)
    figs["distribution"] = make_outlet_distribution()
    return figs


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.78)
    section.bottom_margin = Cm(1.78)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Caption"]:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Heading 1"].font.size = Pt(10)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(10)
    doc.styles["Heading 2"].font.bold = True


def set_columns(section, num: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(num))
    cols_el.set(qn("w:space"), "360")


def add_section(doc: Document, columns: int, start=WD_SECTION_START.CONTINUOUS):
    section = doc.add_section(start)
    section.top_margin = Cm(1.78)
    section.bottom_margin = Cm(1.78)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    set_columns(section, columns)
    return section


def add_para(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size: float = 10.0) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.italic = True


def add_picture(doc: Document, path: Path, caption: str, width_cm: float) -> None:
    doc.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_compact_table(doc: Document, final: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    rows = [
        ("Target correct", f"{f(final, 'mean_target_correct'):.3f} +- {f(final, 'std_target_correct'):.3f}"),
        ("Live recovery", f"{f(final, 'mean_live_recovery'):.3f}"),
        ("Dead removal", f"{f(final, 'mean_dead_removal'):.3f}"),
        ("Wall loss", f"{f(final, 'mean_wall_loss'):.3f}"),
        ("Topology gain", f"{f(final, 'topology_gain_vs_straight'):.3f}"),
        ("Voltage", f"{f(final, 'voltage_v'):.2f} V"),
        ("Flow rate", f"{f(final, 'flow_uL_min'):.3f} uL/min"),
        ("Joule power proxy", f"{f(final, 'active_joule_power_mW'):.2f} mW"),
    ]
    for key, val in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = val
    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(4.2)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(7.2)


def build_ieee_docx(figs: dict[str, Path]) -> None:
    sample, final = final_sample_from_v12()
    doc = Document()
    set_doc_style(doc)
    set_columns(doc.sections[0], 1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Compact Spiral Dielectrophoretic Microfluidic Design for Continuous Live/Dead Cell Sorting")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    add_para(doc, "Analytical Design Project - IEEE-style simulation manuscript draft", WD_ALIGN_PARAGRAPH.CENTER, 10)
    add_para(
        doc,
        "Abstract - This manuscript draft presents a compact spiral dielectrophoretic (DEP) microfluidic design for continuous live/dead cell sorting. The work responds to a specific modeling gap: spiral microchannels and DEP separators are well established separately, and spiral DEP devices exist, but a compact spiral DEP live/dead sorter requires defensible device geometry, mechanism controls, thermal checks, and clear visualization. The final reduced-order V12 operating point reaches 0.940 +/- 0.013 target correct, 0.880 live recovery, 1.000 dead removal, 0.000 wall loss, 0.547 uL/min flow rate, and 2.86 mW active Joule-power proxy. A same-length straight DEP control gives 0.809 target correct, leaving a topology gain of 0.131. The result is framed as a computational design candidate rather than experimental proof.",
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        9.2,
    )
    add_para(doc, "Index Terms - dielectrophoresis, spiral microfluidics, live/dead sorting, numerical design, OpenFOAM-oriented workflow.", WD_ALIGN_PARAGRAPH.JUSTIFY, 9.0)
    add_picture(doc, figs["device"], "Fig. 1. Selected device-level V14 schematic. Unlike earlier centerline-only figures, this panel shows channel ribbon, two side-wall electrode traces, inlet, two outlet branches, and operating dimensions.", 17.0)
    add_picture(doc, figs["library"], "Fig. 2. V14 clean geometry library. Every candidate is drawn as a finite-width device channel with side electrodes and inlet/outlet markers; the CSV file stores formula, length, footprint, and clearance gate.", 17.0)

    add_section(doc, 2, WD_SECTION_START.CONTINUOUS)
    add_heading(doc, "I. Introduction")
    add_para(
        doc,
        "Live/dead cell sorting is important in cell processing, bioreactor monitoring, and label-free sample preparation. Passive spiral microchannels can provide compact curved-channel focusing, but viability contrast is weak when live and dead cells have overlapping size distributions. DEP provides an orthogonal electrical-property contrast through the complex Clausius-Mossotti factor. The design hypothesis is therefore not simply that a longer spiral improves separation, but that a compact curved channel and staged DEP activation can provide useful topology gain beyond a same-length straight DEP channel.",
    )
    add_para(
        doc,
        "The 2026 spiral DEP numerical paper in the local literature set is used as a figure-language reference: clean channel boundaries, electrode labels, inlet/outlet details, and annotated device regions. V14 adopts that style while keeping the project's own result bounded. The current model remains reduced order and must not be read as electrode-resolved CFD/FEM validation.",
    )
    add_heading(doc, "II. Device Geometry and Model")
    add_para(
        doc,
        "The selected topology is a monotone-curvature spiral generated by r(theta)=r0/(1+k theta)^p with a smooth outer entry segment. Side-wall DEP electrodes are abstracted as two active traces along the channel boundaries, with active coverage between the optimized DEP start and end fractions. The outlet is represented as a two-branch splitter that maps lateral position into live-rich and dead-rich outlets.",
    )
    add_para(
        doc,
        f"The final channel length is {path_length(sample)*1e3:.2f} mm. The optimized operating point uses {sample['voltage_v']:.2f} V, {sample['frequency_hz']/1e6:.2f} MHz, {sample['electrode_gap_m']*1e6:.1f} um electrode gap, and {f(final, 'flow_uL_min'):.3f} uL/min flow. Population-level trajectories are sampled from live/dead dielectric and size distributions used in the V12 package.",
    )
    add_heading(doc, "III. Formula-Defined Geometry Library")
    add_para(
        doc,
        "Earlier versions over-emphasized centerlines and visually irregular curves. V14 therefore records a larger formula-defined library using channel ribbons, as shown in Fig. 2: Archimedean, elliptic, monotone-curvature, logarithmic, Fermat, superellipse/oval, C-arc, S-bend-prefocused, contraction-entry, expanded-outlet, D-shaped, and smooth kidney-like variants. These are not claimed as all simulated finalists; they are the clean topology candidates that can be carried into the next high-fidelity mesh stage.",
    )
    add_heading(doc, "IV. Optimization and ML Use")
    add_para(
        doc,
        "ML is used as a surrogate optimizer, not as a black-box claim generator. ExtraTrees regression ranks formula-parameter candidates from reduced-order simulations. The objective penalizes unrealistic apparent maxima by checking same-length straight DEP controls, no-DEP controls, inlet-focus controls, wall loss, thermal proxy, pressure proxy, and geometry clearance. This prevents the trivial route of increasing channel length or voltage until the reduced-order model reports near-perfect separation.",
    )
    add_picture(doc, figs["feature_importance"], "Fig. 3. Existing V12 surrogate feature-importance plot retained for traceability of the ML-assisted optimization stage.", 7.9)
    add_heading(doc, "V. Results")
    add_para(
        doc,
        "The final selected point deliberately sacrifices some apparent maximum accuracy for lower thermal load and stronger device plausibility. It passes the target gate, retains positive topology gain, and keeps wall loss at zero in the held-out validation set. The trajectory overlay in Fig. 4 is generated on the device ribbon rather than on a bare centerline, making the outlet mapping and electrode region easier to inspect.",
    )
    add_picture(doc, figs["trajectory"], "Fig. 4. Representative live/dead trajectory overlay on the selected V14 device schematic. Green and red lines show sampled live and dead paths from the V13 trajectory table mapped onto the final channel ribbon.", 7.9)
    add_picture(doc, figs["results_plate"], "Fig. 5. IEEE-style compact result plate: validation metrics, voltage-performance screen, mechanism controls, and device-operation proxies.", 7.9)
    add_compact_table(doc, final)
    add_picture(doc, figs["distribution"], "Fig. 6. Final lateral-position distribution from the V13 visualization run.", 7.9)
    add_heading(doc, "VI. Thermal and Hydraulic Plausibility")
    add_para(
        doc,
        "The thermal model remains a conservative ranking proxy. Active Joule power is estimated from medium conductivity, field proxy, active volume, and active DEP fraction; the steady substrate delta-T value is not a validated chip temperature. This is sufficient for early-stage school-project screening because it prevents high-voltage over-selection, but a final paper would still require electrode-resolved electric-field and heat-transfer simulation.",
    )
    add_picture(doc, figs["thermal"], "Fig. 7. V12 thermal feasibility map retained as the operating-window safety check for voltage and electrode gap.", 7.9)
    add_heading(doc, "VII. Discussion and Limitations")
    add_para(
        doc,
        "The most important interpretation is negative as well as positive. A high correct value alone is not academically persuasive because the reduced-order model can be made over-deterministic by increasing residence time or field strength. The useful result is the comparison against controls: no DEP is near chance, same-length straight DEP is lower than the selected spiral, and the selected design keeps thermal and wall-loss proxies within plausible ranges.",
    )
    add_para(
        doc,
        "Remaining limitations are substantial: the current electric field is not solved from explicit electrode geometry, thermal transport is simplified, cell deformation and cell-cell interaction are omitted, and dielectric parameters require calibration for a real cell line. Therefore V14 is visually and structurally closer to a manuscript, but the scientific claim should remain: a candidate device and optimization workflow have been identified for high-fidelity validation.",
    )
    add_heading(doc, "VIII. Conclusion")
    add_para(
        doc,
        "V14 produces a cleaner IEEE-style manuscript draft and a more credible device figure package. The final candidate is a compact monotone-curvature spiral DEP sorter with explicit electrodes, inlet, outlet, dimensions, trajectory overlay, and multi-objective operating metrics. The next technical step is to convert the selected ribbon geometry into an electrode-resolved OpenFOAM/FEM case and test whether the reported topology gain survives the higher-fidelity field model.",
    )
    add_heading(doc, "References")
    refs = [
        "[1] S. Betyar and A. Ramiar, Numerical Simulation of Biological Particle Separation in a Spiral Microchannel Using the Dielectrophoresis Mechanism, 2026.",
        "[2] A. Shafiee et al., Contactless dielectrophoresis for selective live/dead cell isolation, Lab on a Chip, 2010.",
        "[3] S. Elitas et al., Dielectrophoretic separation of live and dead monocytes using 3D carbon electrodes, Sensors, 2017.",
        "[4] T. Kwon and S. Choi et al., continuous removal of small nonviable suspended mammalian cells and debris using inertial microfluidics, Lab on a Chip, 2018.",
        "[5] V12 local analytical-design package, final optimization validation summary and controls, 2026.",
    ]
    for ref in refs:
        add_para(doc, ref, WD_ALIGN_PARAGRAPH.LEFT, 8.5)
    doc.save(REPORT_DOCX)


def render_docx() -> tuple[Path | None, list[Path]]:
    render_dir = OUT_DIR / "docx_render"
    render_dir.mkdir(parents=True, exist_ok=True)
    pdf = render_dir / f"{REPORT_DOCX.stem}.pdf"
    try:
        subprocess.run(
            [
                "soffice",
                f"-env:UserInstallation=file:///tmp/lo_profile_v14_{REPORT_DOCX.stem}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(render_dir),
                str(REPORT_DOCX),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        if pdf.exists():
            subprocess.run(["pdftoppm", "-png", str(pdf), str(render_dir / "page")], check=True)
    except (OSError, subprocess.CalledProcessError):
        return None, []
    pages = sorted(render_dir.glob("page-*.png"))
    return pdf if pdf.exists() else None, pages


def write_readme(figs: dict[str, Path], pdf: Path | None, pages: list[Path]) -> None:
    _, final = final_sample_from_v12()
    lines = [
        "# V14 IEEE Publication Package",
        "",
        "V14 repackages the validated V12/V13 reduced-order results with device-level publication figures and an IEEE-style two-column DOCX draft.",
        "",
        "## What changed",
        "",
        "- Replaced centerline-only main figures with finite-width channel ribbons.",
        "- Added side-wall electrode traces, inlet arrow, two-outlet splitter, and dimension inset.",
        "- Added a larger formula-defined geometry library with clearance screening metadata.",
        "- Built a two-column IEEE-style DOCX manuscript draft.",
        "- Kept the scientific claim bounded: this is still a reduced-order numerical design candidate.",
        "",
        "## Final validated operating point",
        "",
        f"- target correct: `{f(final, 'mean_target_correct'):.3f} +/- {f(final, 'std_target_correct'):.3f}`",
        f"- live recovery: `{f(final, 'mean_live_recovery'):.3f}`",
        f"- dead removal: `{f(final, 'mean_dead_removal'):.3f}`",
        f"- wall loss: `{f(final, 'mean_wall_loss'):.3f}`",
        f"- topology gain vs same-length straight DEP: `{f(final, 'topology_gain_vs_straight'):.3f}`",
        f"- flow rate: `{f(final, 'flow_uL_min'):.3f} uL/min`",
        f"- active Joule-power proxy: `{f(final, 'active_joule_power_mW'):.2f} mW`",
        "",
        "## Artifacts",
        "",
        f"- DOCX: `{REPORT_DOCX.relative_to(ROOT)}`",
        f"- rendered PDF: `{pdf.relative_to(ROOT) if pdf else 'not rendered'}`",
        f"- rendered pages: `{len(pages)}`",
        "- geometry library metadata: `results/design_v14_ieee_publication_package/v14_geometry_library.csv`",
        "- figures: `results/design_v14_ieee_publication_package/figures/`",
        "",
        "## Figures",
    ]
    for key, path in figs.items():
        lines.append(f"- `{key}`: `{path.relative_to(ROOT)}`")
    lines.extend(
        [
            "",
            "## Honesty note",
            "",
            "The figures are publication-style, but publication readiness of the science still depends on a higher-fidelity electric-field/thermal validation stage. V14 is a strong manuscript draft, not final experimental proof.",
        ]
    )
    (OUT_DIR / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    ensure_dirs()
    figs = build_figures()
    build_ieee_docx(figs)
    pdf, pages = render_docx()
    write_readme(figs, pdf, pages)
    print(
        json.dumps(
            {
                "docx": str(REPORT_DOCX),
                "out_dir": str(OUT_DIR),
                "figures": {k: str(v) for k, v in figs.items()},
                "rendered_pdf": str(pdf) if pdf else None,
                "rendered_pages": len(pages),
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
