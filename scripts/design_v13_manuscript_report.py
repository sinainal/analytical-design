#!/usr/bin/env python3
"""V13 manuscript-style DOCX report for the spiral DEP design study.

This script does not invent new numerical results. It turns the V9-V12 design
history into a paper-like manuscript and adds publication-style simulation
visuals for the selected V12 operating point.
"""

from __future__ import annotations

import csv
import json
import math
import shutil
import subprocess
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import Cm, Pt, RGBColor

from design_v11_shape_ml_nonintersection import centerline, path_length, sample_to_spec, sigmoid
from study_common import load_openfoam_field_stats, run_openfoam_case, run_particle, sample_cell


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "results" / "design_v13_manuscript_report"
FIG_DIR = OUT_DIR / "figures"
DOC_DIR = ROOT / "output" / "doc"
REPORT_DOCX = DOC_DIR / "spiral_dep_v13_manuscript_draft.docx"


def read_csv(path: Path) -> list[dict]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def f(row: dict, key: str, default: float = 0.0) -> float:
    try:
        return float(row.get(key, default))
    except (TypeError, ValueError):
        return default


def save_plot(fig: plt.Figure, name: str) -> Path:
    path = FIG_DIR / name
    fig.tight_layout()
    fig.savefig(path, dpi=330, bbox_inches="tight")
    plt.close(fig)
    return path


def final_sample_from_v12() -> dict:
    sample = load_json(ROOT / "results" / "design_v11_shape_ml_nonintersection" / "rank1_monotone_curvature_spiral_formula.json")
    best = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_final_validation_summary.csv")[0]
    sample["voltage_v"] = f(best, "voltage_v")
    sample["velocity_m_s"] = f(best, "velocity_um_s") * 1e-6
    sample["electrode_gap_m"] = f(best, "electrode_gap_um") * 1e-6
    sample["electrode_coverage"] = f(best, "electrode_coverage")
    sample["frequency_hz"] = f(best, "frequency_hz")
    return sample


def mapped_path_geometry(sample: dict, path_rows: list[dict]) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    _, xc, yc = centerline(sample, n=1800)
    dx = np.gradient(xc)
    dy = np.gradient(yc)
    norm = np.maximum(np.hypot(dx, dy), 1e-12)
    nx = -dy / norm
    ny = dx / norm
    xout, yout, pout = [], [], []
    for row in path_rows:
        p = float(row["progress"])
        idx = min(len(xc) - 1, max(0, int(round(p * (len(xc) - 1)))))
        lateral = float(row["lateral_m"])
        xout.append(xc[idx] + nx[idx] * lateral)
        yout.append(yc[idx] + ny[idx] * lateral)
        pout.append(p)
    return np.array(xout), np.array(yout), np.array(pout)


def generate_final_trajectory_visuals(sample: dict) -> dict[str, Path]:
    run_openfoam_case()
    field_stats = load_openfoam_field_stats()
    rng = np.random.default_rng(13130)
    spec = sample_to_spec(sample, effective_voltage=True)
    length_m = path_length(sample)
    paths: list[dict] = []
    results: list[dict] = []
    pid = 0
    for cls in ["live", "dead"]:
        for _ in range(28):
            cell = sample_cell(
                rng,
                cls,
                spec.channel_width_m,
                sample["frequency_hz"],
                sample["inlet_offset_ratio"],
                sample["inlet_spread_ratio"],
            )
            result, path = run_particle(
                spec,
                cell,
                field_stats,
                rng,
                pid,
                260,
                sample["frequency_hz"],
                sample["outlet_split_ratio"],
                geometry="spiral",
                dep_enabled=True,
                dep_sign=sample["dep_sign"],
                keep_path=True,
                length_m=length_m,
                dep_start_fraction=sample["dep_start_fraction"],
                dep_end_fraction=sample["dep_end_fraction"],
                dean_scale=sample["dean_scale"],
            )
            for row in path:
                row["class"] = cls
                row["particle_id"] = pid
                row["outlet"] = result["outlet"]
            paths.extend(path)
            results.append(result)
            pid += 1

    write_rows(paths, OUT_DIR / "v13_final_trajectory_paths.csv")
    write_rows(results, OUT_DIR / "v13_final_trajectory_particles.csv")

    plt.style.use("seaborn-v0_8-whitegrid")
    figs: dict[str, Path] = {}
    _, xcenter, ycenter = centerline(sample, n=1800)
    fig, ax = plt.subplots(figsize=(7.6, 6.6))
    ax.plot(xcenter * 1e3, ycenter * 1e3, color="#0f172a", lw=2.0, alpha=0.75)
    for cls, color in [("live", "#2a9d8f"), ("dead", "#e76f51")]:
        ids = sorted({int(r["particle_id"]) for r in paths if r["class"] == cls})
        for particle_id in ids[::2]:
            rows = [r for r in paths if int(r["particle_id"]) == particle_id]
            x, y, _ = mapped_path_geometry(sample, rows)
            ax.plot(x * 1e3, y * 1e3, color=color, alpha=0.35, lw=0.9)
    ax.set_aspect("equal")
    ax.set_xlabel("x (mm)")
    ax.set_ylabel("y (mm)")
    ax.set_title("Final V13 simulated live/dead trajectories")
    ax.plot([], [], color="#2a9d8f", label="live")
    ax.plot([], [], color="#e76f51", label="dead")
    ax.legend(loc="upper right")
    figs["trajectory_map"] = save_plot(fig, "fig_v13_trajectory_map.png")

    fig, ax = plt.subplots(figsize=(7.4, 5.0))
    live = [r for r in results if r["class"] == "live"]
    dead = [r for r in results if r["class"] == "dead"]
    bins = np.linspace(-60, 35, 28)
    ax.hist([float(r["final_lateral_um"]) for r in live], bins=bins, alpha=0.70, color="#2a9d8f", label="live")
    ax.hist([float(r["final_lateral_um"]) for r in dead], bins=bins, alpha=0.70, color="#e76f51", label="dead")
    ax.axvline(-0.5 * spec.channel_width_m * 1e6 + sample["outlet_split_ratio"] * spec.channel_width_m * 1e6, color="#111827", ls="--", lw=1.2, label="outlet split")
    ax.set_xlabel("Final lateral position (um)")
    ax.set_ylabel("Particle count")
    ax.set_title("Final simulated outlet-position distributions")
    ax.legend()
    figs["final_distribution"] = save_plot(fig, "fig_v13_final_distribution.png")

    fig, ax = plt.subplots(figsize=(7.6, 5.2))
    classes = ["live", "dead"]
    outlets = ["inner", "outer", "lost"]
    data = np.array([[sum(r["class"] == cls and r["outlet"] == outlet for r in results) for outlet in outlets] for cls in classes])
    bottom = np.zeros(len(classes))
    colors = ["#457b9d", "#2a9d8f", "#8d99ae"]
    for i, outlet in enumerate(outlets):
        ax.bar(classes, data[:, i], bottom=bottom, color=colors[i], label=outlet)
        bottom += data[:, i]
    ax.set_ylabel("Particle count")
    ax.set_title("Outlet classification from V13 trajectory visualization run")
    ax.legend(ncol=3)
    figs["outlet_counts"] = save_plot(fig, "fig_v13_outlet_counts.png")

    return figs


def write_rows(rows: list[dict], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        return
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()), lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


def copy_existing_figures() -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    sources = {
        "cm_factor": ROOT / "results" / "cm_factor" / "design_v0_cm_factor.png",
        "v8_professional_still": ROOT / "results" / "design_v8_shape_ml_optimization" / "v8_best_professional_still.png",
        "v11_gallery": ROOT / "results" / "design_v11_shape_ml_nonintersection" / "v11_clean_shape_gallery.png",
        "v12_geometry": ROOT / "results" / "design_v12_final_package" / "figures" / "fig01_final_geometry.png",
        "v12_pareto": ROOT / "results" / "design_v12_final_package" / "figures" / "fig05_operating_pareto_power.png",
        "v12_flow": ROOT / "results" / "design_v12_final_package" / "figures" / "fig06_flowrate_tradeoff.png",
        "v12_thermal": ROOT / "results" / "design_v12_final_package" / "figures" / "fig07_thermal_map.png",
        "v12_metrics": ROOT / "results" / "design_v12_final_package" / "figures" / "fig09_final_metrics.png",
        "v12_controls": ROOT / "results" / "design_v12_final_package" / "figures" / "fig10_final_controls.png",
        "v12_workflow": ROOT / "results" / "design_v12_final_package" / "figures" / "fig12_workflow.png",
    }
    out: dict[str, Path] = {}
    for key, src in sources.items():
        dst = FIG_DIR / f"{key}{src.suffix}"
        shutil.copy2(src, dst)
        out[key] = dst
    return out


def make_manuscript_figures() -> dict[str, Path]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    sample = final_sample_from_v12()
    figs = copy_existing_figures()
    figs.update(generate_final_trajectory_visuals(sample))

    v12 = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_final_validation_summary.csv")
    v11 = read_csv(ROOT / "results" / "design_v11_shape_ml_nonintersection" / "v11_validation_summary.csv")
    fig, ax = plt.subplots(figsize=(8.2, 5.2))
    ax.scatter([f(r, "length_mm") for r in v11], [f(r, "topology_gain_vs_straight") for r in v11], s=80, color="#8d99ae", label="V11 geometry candidates")
    ax.scatter([f(r, "length_mm") for r in v12], [f(r, "topology_gain_vs_straight") for r in v12], s=95, color="#2a9d8f", label="V12 operating candidates")
    ax.axhline(0.08, color="#111827", ls="--", lw=1)
    ax.set_xlabel("Channel length (mm)")
    ax.set_ylabel("Gain vs same-length straight DEP")
    ax.set_title("Topology contribution retained after operating optimization")
    ax.legend()
    figs["topology_gain_comparison"] = save_plot(fig, "fig_v13_topology_gain_comparison.png")

    rows = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_operating_screen.csv")
    fig, ax = plt.subplots(figsize=(8.2, 5.2))
    sc = ax.scatter([f(r, "voltage_v") for r in rows], [f(r, "target_correct") for r in rows], c=[f(r, "active_joule_power_mW") for r in rows], cmap="inferno", s=42, alpha=0.84)
    fig.colorbar(sc, ax=ax, label="Active Joule power proxy (mW)")
    ax.axhline(0.90, color="#111827", ls="--", lw=1)
    ax.set_xlabel("Voltage (V)")
    ax.set_ylabel("Target correct")
    ax.set_title("Voltage-performance-thermal trade-off")
    figs["voltage_tradeoff"] = save_plot(fig, "fig_v13_voltage_tradeoff.png")

    return figs


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    for name, size in [("Normal", 10), ("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(15, 23, 42)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(31, 41, 55)


def p(doc: Document, text: str, style: str | None = None) -> None:
    para = doc.add_paragraph(style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(10)


def caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)


def add_fig(doc: Document, path: Path, cap: str, width_cm: float = 15.8) -> None:
    doc.add_picture(str(path), width=Cm(width_cm))
    caption(doc, cap)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_docx(figs: dict[str, Path]) -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("A Compact Monotone-Curvature Spiral DEP Microfluidic Design for Continuous Live/Dead Cell Sorting")
    r.bold = True
    r.font.size = Pt(20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("V13 manuscript-style simulation report").italic = True
    p(doc, "Prepared from the analytical-design repository on 2026-05-26. This document is a manuscript draft based on reduced-order numerical simulation, not an experimental performance claim.")
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    p(doc, "Spiral microchannels and dielectrophoresis (DEP) have independently been used for biological-particle handling, yet the local literature matrix did not identify a continuous spiral DEP architecture specifically targeting viable/nonviable mammalian-cell sorting with OpenFOAM-style reproducible numerical workflow. This work develops a compact monotone-curvature spiral design and evaluates it with a reduced-order live/dead cell model, mechanism controls, topology controls, and operation-level thermal and hydraulic metrics. The final V12 operating point achieves a validated target correct of 0.940 +/- 0.013 with live recovery of 0.880, dead removal of 1.000, live-outlet purity of 1.000, dead-outlet purity of 0.893, zero wall loss in the validation set, flow rate of 0.547 uL/min, active Joule-power proxy of 2.86 mW, and steady substrate temperature-rise proxy of 3.89 C. A same-length straight DEP control gives 0.809 target correct, leaving a topology gain of 0.131. The result supports the design hypothesis that a clean spiral topology can contribute beyond residence time alone, while also highlighting the remaining need for electrode-resolved field simulation and experimental validation.")
    doc.add_paragraph("Keywords: dielectrophoresis; spiral microfluidics; live/dead cell sorting; numerical design; monotone curvature; thermal feasibility")
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    p(doc, "The removal or enrichment of nonviable cells is a recurring problem in cell-processing workflows, bioreactor monitoring, and label-free sample preparation. Spiral microchannels are attractive because they can generate curved-channel inertial effects in compact footprints, and they have already been used for nonviable-cell and debris handling [R4, R8, R16]. However, passive spiral methods are intrinsically limited when viable and nonviable populations have overlapping size distributions, because the separation force is then not primarily viability-selective.")
    p(doc, "DEP provides an orthogonal contrast mechanism because the DEP force depends on the complex permittivity difference between the cell and suspending medium. Live and dead cells can therefore respond differently as membrane conductivity, cytoplasm conductivity, and membrane integrity change [R5-R7, R9, R12]. Non-spiral DEP designs have demonstrated live/dead or viability-selective separation, and recent numerical DEP work emphasizes explicit separation efficiency, purity, voltage, electrode and flow sweeps [R10, R11].")
    p(doc, "The working novelty of this project is the combination of spiral compactness, DEP viability contrast, and continuous outlet-based sorting. Spiral DEP itself is not new [R1-R3], and live/dead DEP is not new [R5-R7, R9]. The gap is a continuous spiral DEP live/dead sorting design with explicit topology controls, thermal checks, and traceable simulation outputs. The manuscript therefore asks a narrow design question: can a clean, non-intersecting spiral geometry improve live/dead DEP sorting beyond a same-length straight DEP channel?")
    add_fig(doc, figs["v12_workflow"], "Figure 1. Computational workflow used to move from literature gap to formula geometry, ML screening, held-out validation and V12 operating-point selection.")
    doc.add_page_break()

    doc.add_heading("2. Theory and Design Rationale", level=1)
    p(doc, "The DEP force on a spherical particle in a nonuniform electric field scales with the particle volume, the medium permittivity, the real component of the Clausius-Mossotti factor, and the gradient of the squared electric field. Under low-Reynolds-number conditions, lateral particle motion can be approximated by balancing DEP, drag, diffusion, Dean-like drift and weak shear-lift terms. The key design objective is to accumulate enough differential lateral displacement before the outlet bifurcation while avoiding high voltage, wall loss, or excessive residence time.")
    p(doc, "The spiral is treated as a compact interaction-length amplifier, but not accepted automatically as beneficial. The same-length straight DEP control is essential: if a spiral candidate and a straight candidate perform equally, then the result is likely a residence-time or DEP-strength effect rather than a topology effect. This logic is why several very high-accuracy V10-V12 candidates were rejected.")
    add_fig(doc, figs["cm_factor"], "Figure 2. Frequency-dependent Clausius-Mossotti model used to motivate viability-selective DEP contrast.")
    add_fig(doc, figs["v12_geometry"], "Figure 3. Final monotone-curvature spiral geometry with smooth DEP activation region.")
    doc.add_page_break()

    doc.add_heading("3. Numerical Methods", level=1)
    p(doc, "The design workflow uses formula-defined channel centerlines, population-level live/dead particle sampling, reduced-order DEP forcing, and Lagrangian particle tracking. The geometry families are not arbitrary hand-drawn curves; each is generated by a named equation such as Archimedean, elliptic Archimedean, logarithmic, Fermat, superellipse/oval, short C-arc, S-bend prefocus, and monotone-curvature spiral forms. Candidate geometries with self-intersection or channel overlap are rejected before final validation.")
    p(doc, "The final selected geometry is the monotone-curvature spiral r(theta)=r0/(1+k theta)^p with an outer low-curvature entry. This family was selected because it remained short, non-intersecting, and retained positive gain relative to same-length straight DEP controls. The model tracks recovery, purity, dead removal, wall loss, pressure-drop proxy, Joule-power proxy, and a steady substrate temperature-rise proxy.")
    p(doc, "ML is used as a surrogate optimizer, not as an unrestricted shape generator. ExtraTrees regression ranks formula-parameter candidates after reduced-order simulations. The feature set includes geometry family, voltage, frequency, velocity, turns, radius, pitch, width, eccentricity, curvature power, inlet length, inlet focusing, DEP start/end, electrode gap, electrode coverage, field gain and Dean scale. Finalists are re-simulated with held-out seeds and mechanism controls.")
    add_fig(doc, figs["v11_gallery"], "Figure 4. Clean non-intersecting V11 geometry family finalists.")
    add_fig(doc, figs["v12_pareto"], "Figure 5. Operating-window Pareto screen balancing classification against Joule-power proxy.")
    doc.add_page_break()

    doc.add_heading("4. Shape Optimization Results", level=1)
    p(doc, "V9 showed that formula-defined shapes could reach high classification, but the visual designs were still too close to arbitrary modulation. V10 replaced those shapes with named clean geometries and explicitly prioritized short channels. V11 added more geometry families and a hard non-intersection/channel-overlap filter. This progression changed the design objective from maximizing correct classification to finding a defensible device geometry.")
    p(doc, "The strongest V11 geometry candidate reached 0.963 +/- 0.004 target correct at 7.08 mm channel length, no wall loss, no self-intersection, minimum nonlocal clearance of 238 um, and topology gain of 0.185. This design then became the baseline for V12 operating-condition optimization.")
    add_fig(doc, figs["topology_gain_comparison"], "Figure 6. Topology gain versus length for V11 and V12 candidates.")
    add_fig(doc, figs["v12_controls"], "Figure 7. Mechanism controls comparing validated candidates with same-length straight DEP and no-DEP cases.")
    doc.add_page_break()

    doc.add_heading("5. Final V12 Operating Optimization", level=1)
    p(doc, "The V12 operating sweep optimized voltage, flow velocity, electrode gap, electrode coverage, frequency, inlet focusing, DEP active window and outlet split around the V11 geometry. The final selected point is not the highest-accuracy point. It is a multi-objective compromise that preserves topology gain while keeping the thermal proxy low.")
    final = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_final_validation_summary.csv")[0]
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Target correct", f"{f(final, 'mean_target_correct'):.3f} +/- {f(final, 'std_target_correct'):.3f}"],
            ["Live recovery", f"{f(final, 'mean_live_recovery'):.3f}"],
            ["Dead removal", f"{f(final, 'mean_dead_removal'):.3f}"],
            ["Live outlet purity", f"{f(final, 'mean_live_outlet_purity'):.3f}"],
            ["Dead outlet purity", f"{f(final, 'mean_dead_outlet_purity'):.3f}"],
            ["Flow rate", f"{f(final, 'flow_uL_min'):.3f} uL/min"],
            ["Voltage", f"{f(final, 'voltage_v'):.2f} V"],
            ["Electrode gap", f"{f(final, 'electrode_gap_um'):.1f} um"],
            ["Power proxy", f"{f(final, 'active_joule_power_mW'):.2f} mW"],
            ["Thermal proxy", f"{f(final, 'steady_substrate_delta_c_proxy'):.2f} C"],
            ["Pressure-drop proxy", f"{f(final, 'pressure_drop_kPa'):.2f} kPa"],
            ["Topology gain", f"{f(final, 'topology_gain_vs_straight'):.3f}"],
        ],
    )
    add_fig(doc, figs["v12_metrics"], "Figure 8. Final recovery, removal and purity metrics.")
    doc.add_page_break()

    doc.add_heading("6. Simulation Visualizations", level=1)
    p(doc, "The following figures are included to make the simulation state visible rather than relying only on tables. The trajectory map shows representative live and dead particle paths projected onto the final centerline. The outlet distributions show how the final lateral positions map into the outlet decision.")
    add_fig(doc, figs["trajectory_map"], "Figure 9. Simulated live/dead trajectories on the final V13 geometry visualization run.")
    add_fig(doc, figs["final_distribution"], "Figure 10. Final lateral position distributions for simulated live and dead particles.")
    doc.add_page_break()
    add_fig(doc, figs["outlet_counts"], "Figure 11. Outlet classification count from the V13 visualization run.")
    add_fig(doc, figs["v8_professional_still"], "Figure 12. Previously generated professional-style simulation still used as a visual reference for trajectory rendering.")
    doc.add_page_break()

    doc.add_heading("7. Thermal and Hydraulic Feasibility", level=1)
    p(doc, "Thermal feasibility is handled conservatively as a ranking proxy rather than an absolute claim. The model estimates active Joule power from field, conductivity, volume and active DEP fraction, then uses a simple substrate heat-loss proxy. This cannot replace a heat-transfer solve, but it prevents high-voltage candidates from being selected solely because they give high classification.")
    p(doc, "The selected V12 point has active Joule-power proxy 2.86 mW, steady substrate temperature-rise proxy 3.89 C, flowrate 0.547 uL/min, and pressure-drop proxy 0.10 kPa. These values place the candidate in a plausible operating region for a school-project numerical design study, while keeping the future need for full thermal validation explicit.")
    add_fig(doc, figs["v12_thermal"], "Figure 13. Thermal feasibility map across voltage and electrode gap.")
    add_fig(doc, figs["v12_flow"], "Figure 14. Flowrate and residence-time trade-off.")
    doc.add_page_break()
    add_fig(doc, figs["voltage_tradeoff"], "Figure 15. Voltage-performance-thermal trade-off for V12 operating candidates.")
    doc.add_page_break()

    doc.add_heading("8. Discussion", level=1)
    p(doc, "The main result is not that the model can produce high correct classification. Earlier screens showed that this would be too easy: increasing residence time or using strong DEP can drive the reduced-order model toward near-perfect classification. The more defensible result is that, after adding same-length straight controls, non-intersection checks, wall-loss tracking and thermal proxies, a short monotone-curvature spiral still retains a positive topology contribution.")
    p(doc, "The selected V12 point trades some accuracy for lower thermal load. This is an academically stronger choice than reporting the maximum apparent correct classification because device design must balance separation, sample throughput, heating, pressure, fabrication constraints and robustness. The final topology gain of 0.131 indicates that spiral geometry is contributing, though it is not the only mechanism: inlet focusing and DEP staging remain essential.")
    p(doc, "The project should therefore be framed as a computational design and prioritization study. It proposes a geometry and operating point for the next fidelity level, rather than claiming experimental proof. The next step is electrode-resolved OpenFOAM/FEM field solving on the exact V12 geometry, followed by a thermal solve and a denser uncertainty analysis.")
    doc.add_page_break()

    doc.add_heading("9. Limitations", level=1)
    limitations = [
        "The DEP forcing still uses an electrode-gap/field-gain proxy rather than a full electrode-resolved 3D field solution.",
        "The thermal model is a ranking proxy and should not be treated as a validated chip temperature prediction.",
        "Cell deformation, cell-cell interaction, fouling dynamics and viability after sorting are not included.",
        "The live/dead dielectric model is based on literature parameter ranges and should be recalibrated for a specific cell line.",
        "The report is manuscript-style, but the study still needs a final high-fidelity simulation section before submission-quality publication.",
    ]
    for item in limitations:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    doc.add_heading("10. Conclusion", level=1)
    p(doc, "This V13 manuscript draft presents a compact monotone-curvature spiral DEP design for continuous live/dead sorting as a numerical design study. The final operating point achieves target correct 0.940 +/- 0.013 while preserving a same-length-straight topology gain of 0.131 and maintaining low thermal proxy values. The design is short, non-intersecting, and supported by mechanism controls, ML-assisted parameter screening, and explicit device metrics. The conclusion is deliberately bounded: the geometry is a strong candidate for the next high-fidelity field and thermal simulation stage, not yet an experimentally validated separator.")
    doc.add_page_break()

    doc.add_heading("References", level=1)
    refs = [
        "R1. Betyar and Ramiar, Numerical Simulation of Biological Particle Separation in a Spiral Microchannel Using the Dielectrophoresis Mechanism, 2026.",
        "R2. Patel and Xuan, curvature-induced DEP in spiral/curved microchannels, 2011.",
        "R3. Yilmaz, Spiral Channel Dielectrophoretic Separator, thesis, 2010.",
        "R4. Kwon/Choi et al., continuous removal of small nonviable suspended mammalian cells and debris from bioreactors using inertial microfluidics, Lab on a Chip, 2018.",
        "R5. Elitas et al., Dielectrophoretic Separation of Live and Dead Monocytes Using 3D Carbon-Electrodes, Sensors, 2017.",
        "R6. Patel et al., reservoir-based DEP live/dead yeast separation, Biomicrofluidics, 2012.",
        "R7. Shafiee et al., selective isolation of live/dead cells using contactless DEP, Lab on a Chip, 2010.",
        "R8. Martel and Toner, inertial focusing in spiral microchannels, 2012.",
        "R9. Huang et al., high-efficiency ODEP live/dead cell separation platform, 2025.",
        "R10. Nguyen et al., facing-electrode DEP cell separation, Scientific Reports, 2024.",
        "R11. Tian et al., on-chip DEP single-cell manipulation review, 2024.",
        "R15. Numerical simulations of combined DEP and AC electrothermal flow, Micromachines, 2024.",
        "R16. Petruzzellis et al., lab-on-chip systems for cell sorting and spiral inertial focusing, Micromachines, 2024.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.save(REPORT_DOCX)


def render_docx() -> None:
    render_dir = OUT_DIR / "docx_render"
    render_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation=file:///tmp/lo_profile_v13_{REPORT_DOCX.stem}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(render_dir),
            str(REPORT_DOCX),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    pdf = render_dir / f"{REPORT_DOCX.stem}.pdf"
    if pdf.exists():
        subprocess.run(["pdftoppm", "-png", str(pdf), str(render_dir / "page")], check=True)


def write_readme(figs: dict[str, Path]) -> None:
    final = read_csv(ROOT / "results" / "design_v12_final_package" / "v12_final_validation_summary.csv")[0]
    lines = [
        "# V13 Manuscript Report",
        "",
        "V13 converts the final optimization package into a manuscript-style DOCX with abstract, introduction, theory, methods, results, discussion, limitations, conclusion, references, and simulation visualizations.",
        "",
        "## Final Numerical Result",
        "",
        f"- target correct: `{f(final, 'mean_target_correct'):.3f} +/- {f(final, 'std_target_correct'):.3f}`",
        f"- topology gain vs same-length straight DEP: `{f(final, 'topology_gain_vs_straight'):.3f}`",
        f"- flowrate: `{f(final, 'flow_uL_min'):.3f} uL/min`",
        f"- active Joule power proxy: `{f(final, 'active_joule_power_mW'):.2f} mW`",
        f"- steady substrate delta-T proxy: `{f(final, 'steady_substrate_delta_c_proxy'):.2f} C`",
        "",
        "## Artifacts",
        "",
        f"- DOCX: `{REPORT_DOCX.relative_to(ROOT)}`",
        "- trajectory tables: `v13_final_trajectory_paths.csv`, `v13_final_trajectory_particles.csv`",
        "- figures: `figures/`",
        "- rendered PDF/pages: `docx_render/`",
        "",
        "## Figures",
    ]
    for name, path in figs.items():
        lines.append(f"- `{name}`: `{path.relative_to(ROOT)}`")
    (OUT_DIR / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figs = make_manuscript_figures()
    build_docx(figs)
    render_docx()
    write_readme(figs)
    print(json.dumps({"docx": str(REPORT_DOCX), "out_dir": str(OUT_DIR), "figures": len(figs)}, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
